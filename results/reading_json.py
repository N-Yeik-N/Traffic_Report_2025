import json
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
import os
from docxcompose.composer import Composer
from docxtpl import DocxTemplate
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def _remove_top_bottom_bordes(cell, topValue = True, bottomValue = True):
    tc_pr = cell._element.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    if topValue:
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'nil')
        tc_borders.append(top)
    
    if bottomValue:
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'nil')
        tc_borders.append(bottom)

    tc_pr.append(tc_borders)

def get_color_by_los(los: str) -> str:
    colores = {
        "A": "00B050", # Verde
        "B": "B5E6A2", # Verde amarillento
        "C": "FFFF99", # Amarillo
        "D": "FFD961", # Naranja
        "E": "EB844B", # Naranja rojizo
        "F": "FF3B3B", # Rojo
    }
    return colores.get(los, "FFFFFF") #Blanco por defecto

dictNames = {
    "HVMAD": "Hora Valle Madrugada",
    "HPMAD": "Hora Punta Madrugada",
    "HPM": "Hora Punta Mañana",
    "HVM": "Hora Valle Mañana",
    "HPT": "Hora Punta Tarde",
    "HVT": "Hora Valle Tarde",
    "HPN": "Hora Punta Noche",
    "HVN": "Hora Valle Noche",
}

def _filling_df(df: pd.DataFrame, data: json, rowDf: int, tipicidad: str, state: str, scenario: str):
    for _, nodeName in enumerate(data["nodes_names"]):
        name = nodeName[0]
        for _, listAcess in data["node_results"][name].items():
            sentidoOrigin = listAcess["Sentido"].split("-")
            nombre = listAcess["Nombre"]+" - "+sentidoOrigin[0]
            volumen = int(float(listAcess["Numero de Vehiculos"]))
            cola_prom = listAcess["Longitud de Cola Prom."]
            cola_max = listAcess["Longitud de Cola Max."]
            paradas = listAcess["Demora en Paradas Prom."]
            demora = listAcess["Demora Promedio"]
            los = listAcess["LOS"]
            df.loc[rowDf] = [name, nombre, volumen, cola_prom, cola_max, paradas, demora, los, tipicidad.capitalize(), state, scenario]
            rowDf += 1

    return rowDf, df

def _filling_df2(df: pd.DataFrame, data: json, rowDf: int, tipicidad: str, state: str, scenario: str):
    for number, dictDatos in data["vehicle_performance"].items():
        simrun = number
        delayavg = dictDatos["DelayAvg"]
        delaystopavg = dictDatos["DelayStopAvg"]
        speedavg = dictDatos["SpeedAvg"]
        stopsavg = dictDatos["StopsAvg"]
        vehact = dictDatos["VehAct"]
        veharr = dictDatos["VehArr"]
        demandlatent = dictDatos["DemandLatent"]
        df.loc[rowDf] = [state, simrun, delayavg, delaystopavg, speedavg, stopsavg, vehact, veharr, demandlatent, tipicidad.capitalize(), scenario]
        rowDf += 1

    return rowDf, df

def _filling_df3(df: pd.DataFrame, data: json, rowDf: int, tipicidad: str, state: str, scenario: str):
    for number, dictDatos in data["pedestrian_performance"].items():
        simrun = number
        densavg = dictDatos["DensAvg"]
        flowavg = dictDatos["FlowAvg"]
        normspeedavg = dictDatos["NormSpeedAvg"]
        speedavg = dictDatos["SpeedAvg"]
        stoptmavg = dictDatos["StopTmAvg"]
        travtmavg = dictDatos["TravTmAvg"]
        df.loc[rowDf] = [state, simrun, densavg, flowavg, normspeedavg, speedavg, stoptmavg, travtmavg, tipicidad.capitalize(), scenario]
        rowDf += 1
        
    return rowDf, df

def _combine_all_docx(filePathMaster, filePathsList, finalPath) -> None:
    number_of_sections = len(filePathsList)
    master = Document(filePathMaster)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document(filePathsList[i])
        composer.append(doc_temp)
    composer.save(finalPath)

def _generate_table_ref(objectResultPath, new_text):
    doc_template = DocxTemplate("./templates/template_tablas4.docx")
    new_table = doc_template.new_subdoc(objectResultPath)
    doc_template.render({
        "texto": new_text,
        "tabla": new_table,
    })
    objectResultPathRef = objectResultPath[:-5] + '_REF.docx'
    doc_template.save(objectResultPathRef)
    return objectResultPathRef

def _align_content(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'

    for i in range(len(table.columns)):
        cell = table.cell(0,i)
        cell_xml_element = cell._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'),'B4C6E7')
        table_cell_properties.append(shade_obj)

NODES_TOTRES = [
    "Nodo",
    "VehDelay\n(Avg,Avg,All)",
    "StopDelay\n(Avg,Avg,All)",
    "QLenMax\n(Avg,Avg)",
    "QLenMax\n(Avg,Total)",
    "Vehs\n(Avg,Total,All)",
    "LOS Value\n(Avg,Total,All)",
]

def _read_folders(subareaPath, folderName):
    contentByFolder = {
            "HPM": {
                "Tipico": None,
                "Atipico": None,
            },
            "HPT": {
                "Tipico": None,
                "Atipico": None,
            },
            "HPN": {
                "Tipico": None,
                "Atipico": None,
            },
        }

    folderPath = os.path.join(subareaPath, folderName)

    if os.path.exists(folderPath):
        for tipicidad in ["Tipico", "Atipico"]:
            tipicidadFolder = os.path.join(folderPath, tipicidad)
            tipicidadContent = os.listdir(tipicidadFolder)
            tipicidadContent = [file for file in tipicidadContent if not file.endswith(".ini") and file in ["HPM", "HPT", "HPN"]]

            for scenario in ["HPM", "HPT", "HPN"]:
                jsonPath = os.path.join(tipicidadFolder, scenario, "table.json")
                contentByFolder[scenario][tipicidad] = jsonPath

    return contentByFolder

def result_nodos(jsonPathActual, jsonPathOutputBase, jsonPathOutputProyectado, scenario, tipicidad, df, rowDf) -> None:
    if jsonPathActual:
        with open(jsonPathActual, 'r') as file:
            data = json.load(file)

    if jsonPathOutputBase:
        with open(jsonPathOutputBase, 'r') as file2:
            data2 = json.load(file2)

    if jsonPathOutputProyectado:
        with open(jsonPathOutputProyectado, 'r') as file3:
            data3 = json.load(file3)

    #Creating dataframe
    if jsonPathActual:
        rowDf, df = _filling_df(df, data, rowDf, tipicidad, "Actual", scenario)
    if jsonPathOutputBase:
        rowDf, df = _filling_df(df, data2, rowDf, tipicidad, "Propuesta Base", scenario)
    if jsonPathOutputProyectado:
        rowDf, df = _filling_df(df, data3, rowDf, tipicidad, "Propuesta Proyectada", scenario)
            
    return rowDf, df

def result_vehicular(jsonPathActual, jsonPathOutputBase, jsonPathOutputProyectado, scenario, tipicidad, df, rowDf) -> None:
    if jsonPathActual:
        with open(jsonPathActual, 'r') as file:
            data = json.load(file)

    if jsonPathOutputBase:
        with open(jsonPathOutputBase, 'r') as file2:
            data2 = json.load(file2)

    if jsonPathOutputProyectado:
        with open(jsonPathOutputProyectado, 'r') as file3:
            data3 = json.load(file3)

    #Creating dataframe
    if jsonPathActual:
        rowDf, df = _filling_df2(df, data, rowDf, tipicidad, "Actual", scenario)
    if jsonPathOutputBase:
        rowDf, df = _filling_df2(df, data2, rowDf, tipicidad, "Propuesta Base", scenario)
    if jsonPathOutputProyectado:
        rowDf, df = _filling_df2(df, data3, rowDf, tipicidad, "Propuesta Proyectada", scenario)

    return rowDf, df

def result_peatonal(jsonPathActual, jsonPathOutputBase, jsonPathOutputProyectado, scenario, tipicidad, df, rowDf) -> None:
    if jsonPathActual:
        with open(jsonPathActual, 'r') as file:
            data = json.load(file)

    if jsonPathOutputBase:
        with open(jsonPathOutputBase, 'r') as file2:
            data2 = json.load(file2)

    if jsonPathOutputProyectado:
        with open(jsonPathOutputProyectado, 'r') as file3:
            data3 = json.load(file3)

    #Creating dataframe
    if jsonPathActual:
        rowDf, df = _filling_df3(df, data, rowDf, tipicidad, "Actual", scenario)
    if jsonPathOutputBase:
        rowDf, df = _filling_df3(df, data2, rowDf, tipicidad, "Propuesta Base", scenario)
    if jsonPathOutputProyectado:
        rowDf, df = _filling_df3(df, data3, rowDf, tipicidad, "Propuesta Proyectada", scenario)

    return rowDf, df

def create_tables_nodos(df: pd.DataFrame, tipicidad: str, scenario: str, subareaPath: str, namesByCode: dict):

    ########################
    # Obtaining dataframes #
    ########################

    doc = Document()
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #Writing headers
    headers = [
        "Inters.\nEscenario", "Nombre de Acceso",
        "Volumen\n(veh/h)", "Long. de Cola Prom.\n(m)", "Long. de Cola Max.\n(m)",
        "Demora por Paradas\n(s/veh)", "Demora\n(s/veh)", "LOS\n(A-F)"
        ]
    
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    codigo = df["Intersección"].unique().tolist()[0]

    dfActual = df[df["State"] == "Actual"].reset_index(drop=True)
    #dfActual['Delay'] = dfActual['Delay'].astype(float)
    dfActual['Delay'] = pd.to_numeric(dfActual['Delay'], errors='coerce')
    dfBase = df[df["State"] == "Propuesta Base"].reset_index(drop=True)
    dfProyectado = df[df["State"] == "Propuesta Proyectada"].reset_index(drop=True)
    df['Delay'] = pd.to_numeric(df['Delay'], errors='coerce')

    #########################
    # Creation of paragraph #
    #########################
    #print(df)
    repeatedNames = df.groupby('Nombre')['State'].nunique()
    repeatedNames = repeatedNames[repeatedNames == 3].index.tolist()
    #print(repeatedNames)
    dfRepeated = df[df['Nombre'].isin(repeatedNames)].reset_index(drop=True)
    #print(dfRepeated)
    maxDelayActual = dfRepeated[(dfRepeated['State'] == 'Actual') & (dfRepeated['Delay'].notna())].sort_values(by='Delay', ascending=False).iloc[0]
    nameMaxDelayActual = maxDelayActual['Nombre']
    nombre, sentido = nameMaxDelayActual.split(" - ")
    valueMaxDelayActual = float(maxDelayActual['Delay'])
    
    dfOthers = dfRepeated[dfRepeated['Nombre'] == nameMaxDelayActual][['State', 'Delay']]
    valueDelayBase = float(dfOthers[dfOthers['State'] == 'Propuesta Base']['Delay'].iloc[0])
    valueDelayProyectado = float(dfOthers[dfOthers['State'] == 'Propuesta Proyectada']['Delay'].iloc[0])

    #Cases of comparison:
    if valueMaxDelayActual > valueDelayBase and valueMaxDelayActual > valueDelayProyectado:
        comparison_txt = "Para la propuesta base y propuesta proyectada se presenta una disminución de la demora debido a COMPLETAR MANUAL"

    elif valueMaxDelayActual > valueDelayBase and valueMaxDelayActual < valueDelayProyectado:
        if valueDelayProyectado < 1.5*valueMaxDelayActual:
            comparison_txt = "Para la propuesta base se presenta una disminución de la demora y para la propuesta proyectada un aumento ligero de la demora debido a COMPLETAR MANUAL"
        elif valueDelayProyectado >= 1.5*valueMaxDelayActual:
            comparison_txt = "Para la propuesta base se presenta una disminución de la demora y para la propuesta proyectada un aumento de la demora debido a COMPLETAR MANUAL"

    elif valueMaxDelayActual < valueDelayBase and valueMaxDelayActual > valueDelayProyectado:
        if valueDelayBase < 1.5*valueMaxDelayActual:
            comparison_txt = "Para la propuesta base se presenta un aumento ligero de la demora y para la propuesta proyectada una disminución de la demora debido a COMPLETAR MANUAL"
        elif valueDelayBase >= 1.5*valueMaxDelayActual:
            comparison_txt = "Para la propuesta base se presenta un aumento la demora y para la propuesta proyectada una disminución de la demora debido a COMPLETAR MANUAL"
    
    elif valueMaxDelayActual < valueDelayBase and valueMaxDelayActual < valueDelayProyectado:
        comparison_txt = "Para la propuesta base y propuesta proyectada se presenta un aumento de la demora debido a COMPLETAR MANUAL"

    else:
        comparison_txt = "BORRAR ESTE TEXTO, LOS DATOS SALEN IGUALES EN TODOS LOS ESCENARIOS."

    if scenario == "HPM": turno = "mañana"
    elif scenario == "HPT": turno = "tarde"
    elif scenario == "HPN": turno = "noche"

    VARIABLES = {
        "nominterseccion": namesByCode[codigo],
        "codinterseccion": codigo,
        "nomacceso": nombre,
        "sentido": sentido,
        "delaymax": f"{float(valueMaxDelayActual):.1f}",
        "comparison_txt": comparison_txt,
        "scenario": turno
    }

    resultsFolder = os.path.join(subareaPath, "Tablas", "Results")
    os.makedirs(resultsFolder, exist_ok=True)
    commentPath = os.path.join(resultsFolder, f"{codigo}_{tipicidad.upper()}_{scenario}.docx")

    docComment = DocxTemplate("./templates/template_lista_nodos.docx")
    docComment.render(VARIABLES)
    docComment.save(commentPath)

    #################################
    # Creation of node result table #
    #################################

    start = 1
    for j in range(dfActual.shape[0]):
        newRow = table.add_row()
        for i, column in enumerate(["Intersección", "Nombre", "Volumen", "QueueAvg", "QueueMax", "StopDelay", "Delay", "LOS"]):
            if i == 0: continue
            if column in ["QueueAvg", "QueueMax"]:
                newRow.cells[i].text = str(round(float(dfActual.loc[j, column])))
            elif column in ["StopDelay", "Delay"]:
                newRow.cells[i].text = str(round(float(dfActual.loc[j, column]), 1))
            else:
                newRow.cells[i].text = str(dfActual.loc[j, column])
            
            if column == "LOS":
                color_hex = get_color_by_los(str(dfActual.loc[j, column]))
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
                newRow.cells[i]._element.get_or_add_tcPr().append(shading_elm)

    end = dfActual.shape[0] + start -1

    table.cell(start, 0).text = dfActual.iloc[0, 0] + "\nActual"
    table.cell(start, 0).merge(table.cell(end, 0))

    for cell in table.rows[start].cells[1:]:
        _remove_top_bottom_bordes(cell, topValue=False)
    for cell in table.rows[end].cells[1:]:
        _remove_top_bottom_bordes(cell, bottomValue=False)
    for selectedRow in [i for i in range(start+1,end+1)]:
        for cell in table.rows[selectedRow].cells[1:]:
            _remove_top_bottom_bordes(cell)

    start = end+1
    for j in range(dfBase.shape[0]):
        newRow = table.add_row()
        for i, column in enumerate(["Intersección", "Nombre", "Volumen", "QueueAvg", "QueueMax", "StopDelay", "Delay", "LOS"]):
            if i == 0: continue
            if column in ["QueueAvg", "QueueMax"]:
                newRow.cells[i].text = str(round(float(dfBase.loc[j, column])))
            elif column in ["StopDelay", "Delay"]:
                newRow.cells[i].text = str(round(float(dfBase.loc[j, column]), 1))
            else:
                newRow.cells[i].text = str(dfBase.loc[j, column])
            if column == "LOS":
                color_hex = get_color_by_los(str(dfBase.loc[j, column]))
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
                newRow.cells[i]._element.get_or_add_tcPr().append(shading_elm)
    end = dfBase.shape[0] + start - 1

    table.cell(start, 0).text = dfBase.iloc[0,0] + "\nPropuesta\nBase"
    table.cell(start, 0).merge(table.cell(end, 0))

    for cell in table.rows[start].cells[1:]:
        _remove_top_bottom_bordes(cell, topValue=False)
    for cell in table.rows[end].cells[1:]:
        _remove_top_bottom_bordes(cell, bottomValue=False)
    for selectedRow in [i for i in range(start+1,end+1)]:
        for cell in table.rows[selectedRow].cells[1:]:
            _remove_top_bottom_bordes(cell)

    start = end+1  
    for j in range(dfProyectado.shape[0]):
        newRow = table.add_row()
        for i, column in enumerate(["Intersección", "Nombre", "Volumen", "QueueAvg", "QueueMax", "StopDelay", "Delay", "LOS"]):
            if i == 0: continue
            if column in ["QueueAvg", "QueueMax"]:
                newRow.cells[i].text = str(round(float(dfProyectado.loc[j, column])))
            elif column in ["StopDelay", "Delay"]:
                newRow.cells[i].text = str(round(float(dfProyectado.loc[j, column]), 1))
            else:
                newRow.cells[i].text = str(dfProyectado.loc[j, column])
            if column == "LOS":
                color_hex = get_color_by_los(str(dfProyectado.loc[j, column]))
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
                newRow.cells[i]._element.get_or_add_tcPr().append(shading_elm)
    end = dfProyectado.shape[0] + start - 1

    table.cell(start, 0).text = dfProyectado.iloc[0,0] + "\nPropuesta\nProyectada"
    table.cell(start, 0).merge(table.cell(end, 0))

    for cell in table.rows[start].cells[1:]:
        _remove_top_bottom_bordes(cell, topValue=False)
    for cell in table.rows[end].cells[1:]:
        _remove_top_bottom_bordes(cell, bottomValue=False)
    for selectedRow in [i for i in range(start+1,end)]:
        for cell in table.rows[selectedRow].cells[1:]:
            _remove_top_bottom_bordes(cell)

    #Aesthetics
    _align_content(table)

    for selected_row in [0]:
        for cell in table.rows[selected_row].cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.bold = True

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for id, x in zip([1,7],[4.5,1.5]):
        for cell in table.columns[id].cells:
            cell.width = Cm(x)

    table.style = 'Table Grid'

    #Saving table
    result_nodo_path = os.path.join(subareaPath, "Tablas", f"nodos_{tipicidad}_{scenario}_{codigo}.docx")
    doc.save(result_nodo_path)
    new_text = f"Resultados de la intersección {codigo} en la {dictNames[scenario].lower()} del día {tipicidad.lower()}"
    resultNode = _generate_table_ref(result_nodo_path, new_text)

    return resultNode

def create_tables_vehicular(df: pd.DataFrame, tipicidad: str, scenario: str, subareaPath: str):
    df = df.reset_index(drop=True)

    #######################
    # Obtaining variables #
    #######################

    dfAvg = df[df['SimRun'] == 'Avg'].reset_index(drop=True)
    SELECTEDDATA = {
        "delaypromactual": round(dfAvg.loc[0]["DelayAvg"],2),
        "speedpromactual": round(dfAvg.loc[0]["SpeedAvg"],2),
        "delayprombase": round(dfAvg.loc[1]["DelayAvg"],2),
        "speedprombase": round(dfAvg.loc[1]["SpeedAvg"],2),
        "delaypromproy": round(dfAvg.loc[2]["DelayAvg"],2),
        "speedpromproy": round(dfAvg.loc[2]["SpeedAvg"],2),
    }

    resultsFolder = os.path.join(subareaPath, "Tablas", "Results")
    os.makedirs(resultsFolder, exist_ok=True)
    commentPath = os.path.join(
        resultsFolder, f"VEHICULAR_{tipicidad.upper()}_{scenario}.docx"
    )

    docComment = DocxTemplate("./templates/template_lista_vehicular.docx")
    docComment.render(SELECTEDDATA)
    docComment.save(commentPath)

    ##################
    # Creating Table #
    ##################

    doc = Document()
    table = doc.add_table(rows = 1, cols = 9)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = 'Table Grid'

    headers = [
        "Escenarios", "Num. Sim.", "Demora\nPromedio", "Demora\nParadas\nPromedio", "Velocidad\nPromedio",
        "Paradas\nPromedio", "Veh.\nAct.", "Veh.\nArr.", "Demanda\nLatente"
    ]
    
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    for j in range(df.shape[0]):                                                                                                                
        newRow = table.add_row()
        for i, column in enumerate([
            "Escenario", "SimRun", "DelayAvg", "DelayStopAvg", "SpeedAvg",
            "StopsAvg", "VehAct", "VehArr", "DemandLatent"
        ]):
            if i == 0: continue
            if column == "SimRun":
                newRow.cells[i].text = str(df.loc[j, column])
            elif column in ["DelayAvg", "DelayStopAvg", "SpeedAvg", "StopsAvg"]:
                newRow.cells[i].text = str(round(float(df.loc[j, column]),2))
            elif column in ["VehAct", "VehArr", "DemandLatent"]:
                newRow.cells[i].text = str(int(float(df.loc[j, column])))
    
    counts = df["Escenario"].value_counts()

    countActual = counts["Actual"]
    countBase = counts["Propuesta Base"]
    countProyectada = counts["Propuesta Proyectada"]

    if  countActual > 0:
        table.cell(1,0).text = "Actual"
        table.cell(1,0).merge(table.cell(12,0))
        
        for cell in table.rows[1].cells[1:]:
            _remove_top_bottom_bordes(cell, topValue=False)
        for cell in table.rows[8].cells[1:]:
            _remove_top_bottom_bordes(cell, bottomValue=False)
        for selectedRow in [i for i in range(2,8)]:
            for cell in table.rows[selectedRow].cells[1:]:
                _remove_top_bottom_bordes(cell)

    if  countBase > 0:
        table.cell(13,0).text = "Propuesta Base"
        table.cell(13,0).merge(table.cell(24,0))
        for cell in table.rows[13].cells[1:]:
            _remove_top_bottom_bordes(cell, topValue=False)
        for cell in table.rows[20].cells[1:]:
            _remove_top_bottom_bordes(cell, bottomValue=False)
        for selectedRow in [i for i in range(14,20)]:
            for cell in table.rows[selectedRow].cells[1:]:
                _remove_top_bottom_bordes(cell)

    if countProyectada > 0:
        table.cell(25,0).text = "Propuesta Proyectada"
        table.cell(25,0).merge(table.cell(36,0))
        for cell in table.rows[25].cells[1:]:
            _remove_top_bottom_bordes(cell, topValue=False)
        for cell in table.rows[32].cells[1:]:
            _remove_top_bottom_bordes(cell, bottomValue=False)
        for selectedRow in [i for i in range(26,32)]:
            for cell in table.rows[selectedRow].cells[1:]:
                _remove_top_bottom_bordes(cell)

    _align_content(table)

    for selected_row in [0]:
        for cell in table.rows[selected_row].cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.bold = True

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                try:
                    run = paragraph.runs[0]
                    run.font.name = 'Arial Narrow'
                    run.font.size = Pt(11)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    pass

    #Saving table
    result_vehicular_path = os.path.join(subareaPath, "Tablas", f"vehicular_{tipicidad}_{scenario}.docx")
    doc.save(result_vehicular_path)
    new_text = f"Rendimiento de vehículos de la red en la {dictNames[scenario].lower()} día {tipicidad.lower()}"
    resultVehicular = _generate_table_ref(result_vehicular_path, new_text)

    return resultVehicular

def create_tables_peatonal(df: pd.DataFrame, tipicidad: str, scenario: str, subareaPath: str) -> str:
    df = df.reset_index(drop=True)

    #######################
    # Creating paragraphs #
    #######################

    dfAvg = df[df["SimRun"] == "Avg"].reset_index(drop=True)
    VARIABLES = {
        "speedavg_actual": round(float(dfAvg.loc[0]["SpeedAvg"]),2),
        #"stoptmavg_actual": round(float(dfAvg.loc[0]["StopTmAvg"]),2),
        "speedavg_propuesto": round(float(dfAvg.loc[1]["SpeedAvg"]),2),
        #"stoptmavg_propuesto": round(float(dfAvg.loc[1]["StopTmAvg"]),2),
        "speedavg_proyectado": round(float(dfAvg.loc[2]["SpeedAvg"]),2),
        #"stoptmavg_proyectado": round(float(dfAvg.loc[2]["StopTmAvg"]),2),
    }

    paragraphDoc = DocxTemplate("./templates/template_lista_peatonal.docx")
    paragraphDoc.render(VARIABLES)
    paragraphPath = os.path.join(subareaPath, "Tablas", "Results", f"PEATONAL_{tipicidad.upper()}_{scenario}.docx")
    paragraphDoc.save(paragraphPath)

    ##################
    # Creating table #
    ##################

    doc = Document()
    table = doc.add_table(rows = 1, cols = 7)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = 'Table Grid'

    headers = [
        "Escenarios", "Num. Sim", "Dens. Prom.", "Flujo Prom.", "Vel. Norm. Prom.", "Vel. Prom.", "Tiempo Viaje Prom."
    ]

    for i, header in enumerate(headers):
        table.cell(0,i).text = header

    for j in range(df.shape[0]):
        newRow = table.add_row()
        for i, column in enumerate([
            "Escenario", "SimRun", "DensAvg", "FlowAvg", "NormSpeedAvg", "SpeedAvg", "TravTmAvg"
        ]):
            if i == 0: continue
            if column == "SimRun":
                newRow.cells[i].text = str(df.loc[j, column])
            elif column in ["DensAvg", "FlowAvg", "NormSpeedAvg", "SpeedAvg", "TravTmAvg"]:
                try:
                    newRow.cells[i].text = f"{float(df.loc[j, column]):.4f}"
                except:
                    newRow.cells[i].text = str(df.loc[j, column])

    counts = df["Escenario"].value_counts()

    countActual = counts["Actual"]
    countBase = counts["Propuesta Base"]
    countProyectada = counts["Propuesta Proyectada"]

    if  countActual > 0:
        table.cell(1,0).text = "Actual"
        table.cell(1,0).merge(table.cell(12,0))

        for cell in table.rows[1].cells[1:]:
            _remove_top_bottom_bordes(cell, topValue=False)
        for cell in table.rows[8].cells[1:]:
            _remove_top_bottom_bordes(cell, bottomValue=False)
        for selectedRow in [i for i in range(2,8)]:
            for cell in table.rows[selectedRow].cells[1:]:
                _remove_top_bottom_bordes(cell)

    if  countBase > 0:
        table.cell(13,0).text = "Propuesta Base"
        table.cell(13,0).merge(table.cell(24,0))

        for cell in table.rows[13].cells[1:]:
            _remove_top_bottom_bordes(cell, topValue=False)
        for cell in table.rows[20].cells[1:]:
            _remove_top_bottom_bordes(cell, bottomValue=False)
        for selectedRow in [i for i in range(14,20)]:
            for cell in table.rows[selectedRow].cells[1:]:
                _remove_top_bottom_bordes(cell)

    if countProyectada > 0:
        table.cell(25,0).text = "Propuesta Proyectada"
        table.cell(25,0).merge(table.cell(36,0))

        for cell in table.rows[25].cells[1:]:
            _remove_top_bottom_bordes(cell, topValue=False)
        for cell in table.rows[32].cells[1:]:
            _remove_top_bottom_bordes(cell, bottomValue=False)
        for selectedRow in [i for i in range(26,32)]:
            for cell in table.rows[selectedRow].cells[1:]:
                _remove_top_bottom_bordes(cell)

    _align_content(table)

    for selected_row in [0]:
        for cell in table.rows[selected_row].cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.bold = True

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                try:
                    run = paragraph.runs[0]
                    run.font.name = 'Arial Narrow'
                    run.font.size = Pt(11)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    pass

    #Saving tableq
    result_peatonal_path = os.path.join(subareaPath, "Tablas", f"peatonal_{tipicidad}_{scenario}.docx")
    doc.save(result_peatonal_path)
    new_text = f"Rendimiento de peatones de la red en la {dictNames[scenario].lower()} día {tipicidad.lower()}"
    resultPeatonal = _generate_table_ref(result_peatonal_path, new_text)

    return resultPeatonal

def generate_results(subareaPath: str) -> list[str]:

    resultFolder = os.path.join(subareaPath, "Tablas", "Results")
    os.makedirs(resultFolder, exist_ok=True)

    #############################################
    # Obtaining dictionary with names and codes #
    #############################################

    subareaID = os.path.split(subareaPath)[1]
    subareaID = int(subareaID[-3:])
    namesByCode = {}

    #Obtaining cleaning names finding codes
    excelPath = "./data/Datos Generales.xlsx"
    dfGenData = pd.read_excel(excelPath, sheet_name='DATOS', header=0, usecols=["Code", "Interseccion", "Sub_Area"])
    dfDatos = dfGenData[dfGenData['Sub_Area'] == subareaID].reset_index()

    for _, row in dfDatos.iterrows():
        code = row['Code']
        name = row['Interseccion']
        namesByCode[code] = name

    #######################
    # Reading .json files #
    #######################

    df = pd.DataFrame(
        columns= [
            "Intersección", "Nombre", "Volumen", "QueueAvg", "QueueMax",
            "StopDelay", "Delay", "LOS", "Tipicidad", "State", "Scenario"
        ]
    )
    
    df2 = pd.DataFrame(
        columns=[
            "Escenario", "SimRun", "DelayAvg", "DelayStopAvg", "SpeedAvg", "StopsAvg", "VehAct", "VehArr", "DemandLatent", "Tipicidad", "Scenario"
        ]
    )

    df3 = pd.DataFrame(
        columns=[
            "Escenario", "SimRun", "DensAvg", "FlowAvg", "NormSpeedAvg", "SpeedAvg", "StopTmAvg", "TravTmAvg", "Tipicidad", "Scenario"
        ]
    )

    #Reading Folders
    actualContent = _read_folders(subareaPath, "Actual")
    outputbaseContent = _read_folders(subareaPath, "Output_Base")
    outputproyectadoContent = _read_folders(subareaPath, "Output_Proyectado")

    rowDf = 0
    rowDf2 = 0
    rowDf3 = 0
    for scenario in ["HPM", "HPT", "HPN"]:
        for tipicidad in ["Tipico", "Atipico"]:
            jsonPathActual = actualContent[scenario][tipicidad]
            jsonPathOutputBase = outputbaseContent[scenario][tipicidad]
            jsonPathOutputProyectado = outputproyectadoContent[scenario][tipicidad]

            rowDf, df = result_nodos(
                jsonPathActual,
                jsonPathOutputBase,
                jsonPathOutputProyectado,
                scenario,
                tipicidad.lower(),
                df,
                rowDf
            )
            
            rowDf2, df2 = result_vehicular(
                jsonPathActual,
                jsonPathOutputBase,
                jsonPathOutputProyectado,
                scenario,
                tipicidad.lower(),
                df2,
                rowDf2
            )

            rowDf3, df3 = result_peatonal(
                jsonPathActual,
                jsonPathOutputBase,
                jsonPathOutputProyectado,
                scenario,
                tipicidad.lower(),
                df3,
                rowDf3
            )

    ########################
    # Resultados por nodos #
    ########################

    results_nodes = {
        "Tipico": {
            "HPM": None,
            "HPT": None,
            "HPN": None,
        },
        "Atipico": {
            "HPM": None,
            "HPT": None,
            "HPN": None,
        },
    }

    paragraphsNodes = {
        "Tipico": {
            "HPM": None,
            "HPT": None,
            "HPN": None,
        },
        "Atipico": {
            "HPM": None,
            "HPT": None,
            "HPN": None,
        },
    }

    intersecciones = df['Intersección'].unique().tolist()
    for tipicidad in ["Tipico", "Atipico"]:
        for turno in ["HPM", "HPT", "HPN"]:
            listPaths = []    
            for ints in intersecciones:
                filtered_df = df[
                    (df["Intersección"] == ints) & 
                    (df["Scenario"] == turno) &
                    (df["Tipicidad"] == tipicidad)
                    ]
                #print(filtered_df)# losn nombres de los links debe ser igal en hpm, hpt y hpn, o el nombre de la carpeta esta mal
                tablaPath = create_tables_nodos(filtered_df, tipicidad, turno, subareaPath, namesByCode)
                listPaths.append(tablaPath)

            # Finding list of paragraphs by peak hour #
            paragraphNodeList = []
            for ints in intersecciones:
                paragraphNodeList.append(
                    os.path.join(
                        resultFolder, f"{ints}_{tipicidad.upper()}_{turno}.docx"
                    )
                )
            
            paragraphNodePath = os.path.join(subareaPath, "Tablas", "Results", f"NodesList_{tipicidad.upper()}_{turno}.docx")
            filePathMaster = paragraphNodeList[0]
            filePathList = paragraphNodeList[1:]
            _combine_all_docx(filePathMaster, filePathList, paragraphNodePath)

            paragraphsNodes[tipicidad][turno] = paragraphNodePath
            
            nodoPath = os.path.join(subareaPath, "Tablas", f"nodeResults_{tipicidad}_{turno}_REF.docx")
            if len(listPaths) > 1:
                filePathMaster = listPaths[0]
                filePathList = listPaths[1:]
                _combine_all_docx(filePathMaster, filePathList, nodoPath)
            else:
                nodoPath = listPaths[0]

            results_nodes[tipicidad][turno] = nodoPath
            

    ########################################
    # Resultados por rendimiento vehicular #
    ########################################

    results_vehicular = {
        "Tipico": {
            "HPM": None,
            "HPT": None,
            "HPN": None,
        },
        "Atipico": {
            "HPM": None,
            "HPT": None,
            "HPN": None,
        },
    }

    paragraphsVehicular = {
        "Tipico": {
            "HPM": None,
            "HPT": None,
            "HPN": None,
        },
        "Atipico": {
            "HPM": None,
            "HPT": None,
            "HPN": None,
        },
    }

    for tipicidad in ["Tipico", "Atipico"]:
        for turno in ["HPM", "HPT", "HPN"]:
            filtered_df = df2[
                (df2["Scenario"] == turno) &
                (df2["Tipicidad"] == tipicidad)
                ]
    
            vehicularResultPathRef = create_tables_vehicular(filtered_df, tipicidad, turno, subareaPath)
            results_vehicular[tipicidad][turno] = vehicularResultPathRef

            paragraphsVehicular[tipicidad][turno] = os.path.join(resultFolder, f"VEHICULAR_{tipicidad.upper()}_{turno}.docx")
            
    #######################################
    # Resultados por rendimiento peatonal #
    #######################################

    results_peatonal = {
        "Tipico": {
            "HPM": None,
            "HPT": None,
            "HPN": None,
        },
        "Atipico": {
            "HPM": None,
            "HPT": None,
            "HPN": None,
        },
    }

    paragraphsPeatonal = {
        "Tipico": {
            "HPM": None,
            "HPT": None,
            "HPN": None,
        },
        "Atipico": {
            "HPM": None,
            "HPT": None,
            "HPN": None,
        },
    }

    for tipicidad in ["Tipico", "Atipico"]:
        for turno in ["HPM", "HPT", "HPN"]:
            filtered_df = df3[
                (df3["Scenario"] == turno) &
                (df3["Tipicidad"] == tipicidad)
                ]
    
            peatonalResultPathRef = create_tables_peatonal(filtered_df, tipicidad, turno, subareaPath)
            results_peatonal[tipicidad][turno] = peatonalResultPathRef

            paragraphsPeatonal[tipicidad][turno] = os.path.join(resultFolder, f"PEATONAL_{tipicidad.upper()}_{turno}.docx")

    return results_nodes, results_vehicular, results_peatonal, paragraphsNodes, paragraphsVehicular, paragraphsPeatonal