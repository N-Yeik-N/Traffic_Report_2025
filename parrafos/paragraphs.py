import os

#docx
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def _font_n_height(paragraph):
    paragraph_run = paragraph.runs[0]
    paragraph_run.font.name = 'Arial Narrow'
    paragraph_run.font.size = Pt(11)
    r = paragraph_run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Narrow')

def cambios_variable(path_subarea, codintersecciones):
    tablasPath = os.path.join(path_subarea, "Tablas")
    
    doc = Document()
    doc.add_heading()
    for code in codintersecciones:
        texto = doc.add_paragraph(f"- Intersección {code}:")
        para = doc.add_paragraph("COLOCAR CAMBIOS")
        para.style = 'List Bullet'
        _font_n_height(texto)
        _font_n_height(para)

    finalPath = os.path.join(tablasPath, "cambios.docx")
    doc.save(finalPath)

    return finalPath