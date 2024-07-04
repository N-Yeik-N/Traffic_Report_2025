import os
import pandas as pd
from pathlib import Path
from tables.tools.traffic_lights import get_info
from tables.tools.cycles import get_dates_cycles

#docx
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_table9(path_subarea):
    path_parts = path_subarea.split("/")
    subarea_id = path_parts[-1]
    proyect_folder = '/'.join(path_parts[:-2])

    field_data = os.path.join(
        proyect_folder,
        "7. Informacion de Campo",
        subarea_id,
        "Tiempo de Ciclo Semaforico"
    )

    list_excels = os.listdir(field_data)
    list_excels = [os.path.join(field_data, file) for file in list_excels
                if file.endswith(".xlsx") and not file.startswith("~")]

    phasesList = []
    for excel in list_excels:
        phasesData = get_info(excel)
        phasesList.append(phasesData)

    doc = Document()

    #Típico table
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, txt in enumerate(["Código", "Intersección", "Turno", "T.C.\n(s)", "Fase", "Verde\n(s)", "Ámbar\n(s)", "Rojo\n(s)"]):
        table.cell(0,i).text = txt

    turn = ["HPM", "HPT", "HPN"]
    codeList = []

    rowCount = 1
    for i, excelData in enumerate(phasesList):
        listTipico = excelData.phasesData[:3]
        listCycleTipico = excelData.cycleTimeData[:3]
        startRow = rowCount
        checkFirst = True
        turnCount = 0
        for turnList, cycleTime in zip(listTipico, listCycleTipico): #turnList = [ [120,3,2], [50,3,2]]
            startRowCycle = rowCount
            for faseID, (verde, ambar, rojo) in enumerate(turnList): #[120, 3, 2]
                newRow = table.add_row().cells
                if faseID == 0:
                    if checkFirst == True:
                        newRow[0].text = excelData.codigo
                        codeList.append(excelData.codigo)
                        newRow[1].text = excelData.nombre
                        checkFirst = False
                    newRow[2].text = turn[turnCount]
                    turnCount += 1
                    if turnCount == 3: turnCount = 0
                    newRow[3].text = str(cycleTime)

                newRow[4].text = str(faseID+1)
                newRow[5].text = str(verde)
                newRow[6].text = str(ambar)
                newRow[7].text = str(rojo)
                rowCount += 1

            endRowCycle = rowCount-1
            table.cell(startRowCycle, 2).merge(table.cell(endRowCycle, 2))
            table.cell(startRowCycle, 3).merge(table.cell(endRowCycle, 3))
        endRow = rowCount-1
        table.cell(startRow, 0).merge(table.cell(endRow, 0))
        table.cell(startRow, 1).merge(table.cell(endRow, 1))
        #table.cell(startRow, 2).merge(table.cell(endRow, 2))

    #ESTÉTICA

    for selected_row in [0]:
        for cell in table.rows[selected_row].cells:
            for paragraph in cell.paragraphs:
                try:
                    run = paragraph.runs[0]
                    run.font.bold = True
                except Exception as e: continue

    for i in range(len(table.columns)):
        cell_xml_element = table.rows[0].cells[i]._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'),'B4C6E7')
        table_cell_properties.append(shade_obj)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                try:
                    run = paragraph.runs[0]
                    run.font.name = 'Arial Narrow'
                    run.font.size = Pt(11)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    pass

    for id, x in zip([0,1,4,5,6,7],[0.5,2,0.7,0.5,0.5,0.5]):
        for cell in table.columns[id].cells:
            cell.width = Inches(x)

    table.style = "Table Grid"

    table9_path = Path(path_subarea) / "Tablas" / "table9_sinREF.docx"
    doc.save(table9_path)

    doc_template = DocxTemplate(r"templates\template_tablas.docx")
    if len(codeList) == 1:
        texto = f"Tiempos de ciclo de la intersección {codeList[0]}"
    elif len(codeList) > 1:
        texto_aux = ""
        for i, code in enumerate(codeList):
            if i == len(codeList) - 1:
                texto_aux += f"y {code}"
            elif i == len(codeList) - 2:
                texto_aux += f"{code}"
            else:
                texto_aux += f"{code}, "
            
        texto = f"Tiempos de ciclos y fases de las intersecciones {texto_aux}"

    table9 = doc_template.new_subdoc(table9_path)

    doc_template.render({
        "texto": texto,
        "tabla": table9
    })

    finalPath = os.path.join(path_subarea, "Tablas", "table9.docx")
    doc_template.save(finalPath)

    return finalPath

def create_table8(path_subarea):
    path_excel = r"data\Cronograma Vissim.xlsx"
    df = pd.read_excel(path_excel, "Cronograma", header=0, usecols="A:E", skiprows=1)
    df = df.drop(columns=["Codigo TDR", "Entregable"])

    numsubarea = os.path.split(path_subarea)[1][-3:]
    no = int(numsubarea)

    code_by_subarea = df[df['Sub Area'] == no]['Codigo'].tolist()
    
    try:
        df_dates = get_dates_cycles(code_by_subarea)
    except Exception as e:
        print("Posiblemente no existe un elemento en la base de datos: Datos de Ciclos.xlsx")
        print("Error: ", e)

    doc = Document()
    table = doc.add_table(rows = 7, cols = 5)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #Headers
    for i, header in enumerate(["Intersección", "Día", "Tipicidad", "Turno", "Horario"]):
        table.cell(0,i).text = header

    typicalDay = df_dates['Dia Tipico'].unique().tolist()[0]
    atypicalDay = df_dates['Dia Atipico'].unique().tolist()[0]
    
    for i, day in enumerate([typicalDay, atypicalDay]):
        if i == 0:
            for j in range(1,4,1):
                table.cell(j,1).text = day
        else:
            for j in range(4,7,1):
                table.cell(j,1).text = day

    for i, turno in enumerate(["Mañana", "Tarde", "Noche"]*2):
        table.cell(i+1,3).text = turno

    for i, tipicidad in enumerate(["Típico", "Típico", "Típico", "Atípico", "Atípico", "Atípico"]):	
        table.cell(i+1,2).text = tipicidad

    for i, horario in enumerate([
        "06:30 - 09:30",
        "12:00 - 15:00",
        "17:30 - 20:30",
        "06:30 - 09:30",
        "12:00 - 15:00",
        "17:30 - 20:30",
    ]):
        table.cell(i+1, 4).text = horario

    texto = ""
    for i, code in enumerate(code_by_subarea):
        if i == len(code_by_subarea)-1:
            texto += code
        else:
            texto += code + '\n'

    table.cell(1,0).text = texto
    table.cell(1,0).merge(table.cell(6,0))

    for selected_row in [0]:
        for cell in table.rows[selected_row].cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.bold = True

    for i in range(len(table.columns)):
        cell_xml_element = table.rows[0].cells[i]._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'),'B4C6E7')
        table_cell_properties.append(shade_obj)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for id, x in zip([0,1,2,3,4],
                     [0.9,0.8,0.8,0.5,1]):
        for cell in table.columns[id].cells:
            cell.width = Inches(x)

    table.style = 'Table Grid'
    path_table8 = Path(path_subarea) / "Tablas" / "table8.docx"
    doc.save(path_table8)

    return path_table8