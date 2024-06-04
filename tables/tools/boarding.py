import pandas as pd
from openpyxl import load_workbook
from dataclasses import dataclass
import statistics
from pathlib import Path

#docx
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

@dataclass
class BoardingTable:
    codigo: str
    turno: str
    sentido: str
    acceso: str
    maximo: str
    promedio: str
    desviacion: str

def board_by_excel(path) -> tuple[str, list]:
    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb['Base Data']
    codigo = ws['C4'].value
    #fecha = ws['C5'].value
    columnas = ["Turno", "Sentido", "Acceso", "Tipo de Vehiculo", "Tiempo"]
    df = pd.DataFrame([[cell.value for cell in row] for row in ws['C8:G157']], columns= columnas)
    wb.close()

    turnos = df["Turno"].unique().tolist()
    sentidos = df["Sentido"].unique().tolist()
    accesos = df["Acceso"].unique().tolist()

    tableList = []
    for turno in turnos:
        for sentido in sentidos:
            for acceso in accesos:
                timeSerie = df[(df["Turno"] == turno) & (df["Sentido"] == sentido) & (df["Acceso"] == acceso)]["Tiempo"].tolist()
                if timeSerie == []: continue
                maximo = str(max(timeSerie))
                promedio = str(round(statistics.mean(timeSerie),2))
                if len(timeSerie) == 1:
                    desviacion = "0.0"
                else:
                    desviacion = str(round(statistics.stdev(timeSerie),2))

                info = BoardingTable(
                    codigo = codigo,
                    turno = turno,
                    sentido = sentido,
                    acceso = acceso,
                    maximo = maximo,
                    promedio = promedio,
                    desviacion = desviacion
                )

                tableList.append(info)

    return codigo, tableList

def create_table(
        tableData,  #Lista de un dataclass con la información por fila.
        code,       #Código de intersección
        count,      #Nro. de orden de la tabla
        path_subarea, #Path de la subarea
        ) -> None:
    doc = Document()
    table = doc.add_table(rows=len(tableData)+1, cols=6)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, texto in enumerate(["Turno","Sentido","Acceso", "Máx. de Tiempo", "Prom. Tiempo", "Des. St. Tiempo"]):
        table.cell(0,i).text = texto
    
    row_no = 1
    check = True
    for data in tableData:
        if data.turno == "Mañana":
            if check:
                table.cell(row_no, 0).text = "Mañana"
                check = False
            table.cell(row_no, 1).text = data.sentido
            table.cell(row_no, 2).text = data.acceso
            table.cell(row_no, 3).text = data.maximo
            table.cell(row_no, 4).text = data.promedio
            table.cell(row_no, 5).text = data.desviacion
            row_no += 1

    last_row1 = row_no-1
    table.cell(1,0).merge(table.cell(last_row1,0))
    
    check = True
    for data in tableData:
        if data.turno == "Medio dia":
            if check:
                table.cell(row_no, 0).text = "Medio dia"
                check = False
            table.cell(row_no, 1).text = data.sentido
            table.cell(row_no, 2).text = data.acceso
            table.cell(row_no, 3).text = data.maximo
            table.cell(row_no, 4).text = data.promedio
            table.cell(row_no, 5).text = data.desviacion
            row_no += 1

    last_row2 = row_no-1
    table.cell(last_row1+1,0).merge(table.cell(last_row2,0))

    check = True
    for data in tableData:
        if data.turno == "Noche":
            if check:
                table.cell(row_no, 0).text = "Noche"
                check = False
            table.cell(row_no, 1).text = data.sentido
            table.cell(row_no, 2).text = data.acceso
            table.cell(row_no, 3).text = data.maximo
            table.cell(row_no, 4).text = data.promedio
            table.cell(row_no, 5).text = data.desviacion
            row_no += 1

    last_row3 = row_no-1
    table.cell(last_row2+1,0).merge(table.cell(last_row3,0))

    for selected_row in [0]:
        for cell in table.rows[selected_row].cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.bold = True

    for i in range(len(table.columns)):
        cell_xml_element = table.rows[0].cells[i]._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'),'B4C6E7')
        table_cell_properties.append(shade_obj)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for id, x in zip([0],[0.5]):
        for cell in table.columns[id].cells:
            cell.width = Inches(x)

    table.style = "Table Grid"
    table7_path = Path(path_subarea) / "Tablas" / f"table7_{count}.docx"
    doc.save(table7_path)

    doc_template = DocxTemplate("./templates/template_tablas.docx")
    new_table = doc_template.new_subdoc(table7_path)
    texto = f"Embarque y desembarque de la intersección {code} día típico"
    VARIABLES = {
        'texto': texto,
        'tabla': new_table,
    }
    doc_template.render(VARIABLES)
    table7_path_ref = Path(path_subarea) / "Tablas" / f"table7_{count}_REF.docx"
    doc_template.save(table7_path_ref)
    return table7_path_ref