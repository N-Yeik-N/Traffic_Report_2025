import pandas as pd
import os
from pathlib import Path

#docx
from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_table6(path_subarea):
    path_excel = r"data\Cronograma Vissim.xlsx"
    df = pd.read_excel(path_excel, sheet_name='Cronograma', header=0, usecols="A:E", skiprows=1)
    df = df.drop(columns=["Codigo TDR", "Entregable"])

    numsubarea = os.path.split(path_subarea)[1][-3:]
    no = int(numsubarea)

    code_by_subarea = df[df['Sub Area'] == no]['Codigo'].tolist()
    name_by_subarea = df[df['Sub Area'] == no]['Intersección'].tolist()

    path_eyb = r"data\Paraderos.xlsx"
    df_board = pd.read_excel(path_eyb, sheet_name='Hoja1', header=0, usecols="A:D")
    
    demt_list = []
    dema_list = []
    for code in code_by_subarea:
        demt = df_board[df_board['Código'] == code]['Día Típico'].unique()[0]
        demt = demt.strftime('%d/%m/%Y')
        demt_list.append(demt)
        dema = df_board[df_board['Código'] == code]['Día Atípico'].unique()[0]
        dema = dema.strftime('%d/%m/%Y')
        dema_list.append(dema)
    
    ####################
    # CREATING TABLE 6 #
    ####################

    doc = Document()
    table = doc.add_table(rows = 7, cols = 5)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #Default texts
    table.cell(0,0).text = "Intersección"
    table.cell(0,1).text = "Día"
    table.cell(0,2).text = "Tipicidad"
    table.cell(0,3).text = "Turno"
    table.cell(0,4).text = "Horario"

    #codinterseccion:
    set_codes = ""
    for i, code in enumerate(code_by_subarea):
        if i == len(code_by_subarea)-1:
            set_codes += code
        else:
            set_codes += code + "\n"

    table.cell(1,0).text = set_codes
    table.cell(1,0).merge(table.cell(6,0))

    for i in range(3):
        table.cell(1+i,2).text = "Típico"
        table.cell(4+i,2).text = "Atípico"

    for i, valor in enumerate(["Mañana", "Tarde", "Noche"]):
        table.cell(1+i,3).text = valor
        table.cell(4+i,3).text = valor

    for i, valor in enumerate(["06:30 - 09:30",
                "12:00 - 15:00",
                "17:30 - 20:30",
                "06:30 - 09:30",
                "12:00 - 15:00",
                "17:30 - 20:30",]):
        table.cell(1+i,4).text = valor

    demt = list(set(demt_list))[0]
    dema = list(set(dema_list))[0]

    for i in range(3):
        table.cell(1+i,1).text = demt
        table.cell(4+i,1).text = dema

    for selected_row in [0]:
        for cell in table.rows[selected_row].cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.bold = True

    for i in range(len(table.columns)):
        cell_xml_element = table.rows[0].cells[i]._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'),'B4C6E7')
        table_cell_properties.append(shade_obj)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for id, x in zip([0,1,2,3,4],[0.5,1,1,0.8,1.2]):
        for cell in table.columns[id].cells:
            cell.width = Inches(x)

    table.style = 'Table Grid'
    table6_path = Path(path_subarea) / "Tablas" / "table6.docx"
    doc.save(table6_path)

    doc_template = DocxTemplate("./templates/template_tablas.docx")
    texto = f"Fecha y hora de la recolección de datos de embarque y desembarque"
    new_table = doc_template.new_subdoc(table6_path)

    doc_template.render({"texto": texto, "tabla": new_table})

    finalPath = os.path.join(path_subarea, "Tablas", "boardinTable.docx")
    doc_template.save(finalPath)

    return finalPath