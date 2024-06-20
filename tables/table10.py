from openpyxl import load_workbook
import os
import re
import pandas as pd
#docx
from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _append_document_content(source_doc, target_doc) -> None:
    for element in source_doc.element.body:
        target_doc.element.body.append(element)

def create_table10(path_subarea):
    excelPath = os.path.join(path_subarea, "Program_Results.xlsx")
    if not os.path.exists(excelPath):
        print("Tabla 10\tERROR\tNo existe 'Program_Results.xlsx'")

    #Getting sheets names
    workbook = load_workbook(excelPath, data_only=True, read_only=True)
    sheetNames = workbook.sheetnames
    workbook.close()

    #Selecting sheets
    pattern = r"([A-Z]+-[0-9]+)"
    sheetNames = [name for name in sheetNames if re.match(pattern, name)]

    dict_by_code = {}
    for code in sheetNames:
        df = pd.read_excel(excelPath, code, header=0, usecols="V:AJ", nrows=13)
        dict_by_code[code] = df

    listFinalPathRef = []
    for code, df in dict_by_code.items():
        #Adjust number of columns of the data frame
        firstRow = df.iloc[0].tolist()
        ternList = [firstRow[i:i+3] for i in range(0, len(firstRow),3)]
        noPhases = 5
        for i, tern in enumerate(ternList):
            if sum(tern) == 0:
                noPhases = i
                break
        df = df.drop(df.columns[noPhases*3:], axis=1)
        
        #Creating individual tables
        doc = Document()
        table = doc.add_table(rows = 15, cols = 4+noPhases*3)

        for i, header in enumerate(["Tipicidad","Turno","Hora","TC"]):
            table.cell(0,i).text = header
            table.cell(0,i).merge(table.cell(1,i))

        noAccess = 1
        for no in range(noPhases):
            table.cell(0, 4+no*3).text = f"Acceso {str(noAccess).zfill(2)}-{str(noAccess+1).zfill(2)}"
            table.cell(0, 4+no*3).merge(table.cell(0, 4+no*3+2))
            noAccess += 1

        for i, texto in enumerate(["V", "A", "RR"]*noPhases):
            table.cell(1, 4+i).text = texto

        for i, turno in enumerate(["HVMAD","HPMAD","HPM","HVM","HPT","HVT","HPN","HVN",
                                   "HVMAD","HPM","HPT","HPN","HVN"]):
            table.cell(2+i, 1).text = turno
        
        for i, hora in enumerate(["00:00-05:00", "05:00-06:30", "06:30-10:30", "10:30-12:30", "12:30-15:00", "15:00-17:00", "17:00-22:00", "22:00-00:00",
                                  "00:00-06:00", "06:00-12:00", "12:00-17:00", "17:00-22:00", "22:00-00:00"]):
            table.cell(2+i, 2).text = hora

        for i, fila in df.iterrows():
            table.cell(2+i, 3).text = str(sum(fila))
            for j, elem in enumerate(fila):
                table.cell(2+i, 4+j).text = str(elem)

        table.cell(2, 0).text = "Típico"
        table.cell(2, 0).merge(table.cell(9, 0))

        table.cell(10, 0).text = "Atípico"
        table.cell(10, 0).merge(table.cell(14, 0))

        #Aesthetic:
        for selected_row in [0, 1]:
            for cell in table.rows[selected_row].cells:
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    run.font.bold = True

        for j in [0,1]:
            for i in range(len(table.columns)):
                cell_xml_element = table.rows[j].cells[i]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement('w:shd')
                shade_obj.set(qn('w:fill'),'B4C6E7')
                table_cell_properties.append(shade_obj)

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    run.font.name = 'Arial Narrow'
                    run.font.size = Pt(11)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i != 2: continue
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    run.font.name = 'Arial Narrow'
                    run.font.size = Pt(10)
                    
        listCols = [i for i in range(4+noPhases*3)]
        listCols = listCols[4:]
        listCols[:0] = [1,2,3]
        listWidths = [0.5 for _ in range(4+noPhases*3)][4:]
        listWidths[:0] = [0.5,1.0,0.5]

        for id, x in zip(listCols, listWidths):
            for cell in table.columns[id].cells:
                cell.width = Inches(x)

        table.style = 'Table Grid'

        finalPath = os.path.join(path_subarea, "Tablas", f"tabla10_{code}.docx")
        doc.save(finalPath)

        doc_template = DocxTemplate("./templates/template_tablas2.docx")
        new_table = doc_template.new_subdoc(finalPath) 
        texto = f"Tiempos de ciclos semafóricos en cada escenario para la intersección {code}"
        VARIABLES = {
            'texto': texto,
            'tabla': new_table
        }
        doc_template.render(VARIABLES)
        finalPathRef = os.path.join(path_subarea, "Tablas", f"tabla10_{code}_REF.docx")
        listFinalPathRef.append(finalPathRef)
        doc_template.save(finalPathRef)
    
    #Join tables
    table10_path = os.path.join(path_subarea, "Tablas", "table10.docx")
    doc_target = Document(listFinalPathRef[0])
    for i in range(len(listFinalPathRef)):
        if i == 0: continue
        doc_source = Document(listFinalPathRef[i])
        _append_document_content(doc_source, doc_target)
        doc_target.save(table10_path)
        doc_target = Document(table10_path)

    doc_target.save(table10_path)
    return table10_path