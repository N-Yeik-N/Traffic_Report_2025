import os
import pandas as pd
from tables.tools.pedestrian import *

#docx
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_table12(path_subarea):
    path_excel = r"data\Cronograma Vissim.xlsx"
    df = pd.read_excel(path_excel, "Cronograma", header=0, usecols="A:E", skiprows=1)
    df = df.drop(columns=["Codigo TDR", "Entregable"])

    numsubarea = os.path.split(path_subarea)[1][-3:]
    no = int(numsubarea)

    code_by_subarea = df[df['Sub Area'] == no]['Codigo'].tolist()
    code_n_name = []
    for code in code_by_subarea:
        name_by_code = df[df['Codigo'] == code]['Intersección'].unique()[0]
        code_n_name.append([code, name_by_code])

    doc = Document()
    numberCodes = len(code_by_subarea)
    if numberCodes%2 != 0:
        rowNumbers = numberCodes + 1
    else:
        rowNumbers = numberCodes

    table = doc.add_table(rows=rowNumbers, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (code, name) in enumerate(code_n_name):
        if i%2 == 0: 
            table.cell(i+1, 0).text = f"Intersección {code}\n{name}"
        else:
            table.cell(i, 1).text = f"Intersección {code}\n{name}"

    for row in table.rows:
            for i, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    try:
                        run = paragraph.runs[0]
                        run.font.name = 'Arial Narrow'
                        run.font.size = Pt(11)
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except IndexError:
                         pass

    table.style = "Table Grid"
    finalPath = os.path.join(path_subarea, "Tablas", "table12.docx")
    doc.save(finalPath)

    return finalPath