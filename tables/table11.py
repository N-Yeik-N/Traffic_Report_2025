import os
import pandas as pd
from tables.tools.pedestrian import *

#docx
from docxcompose.composer import Composer
from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _combine_all_docx(filePathMaster, filePathsList, finalPath) -> None:
    number_of_sections = len(filePathsList)
    master = Document(filePathMaster)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document(filePathsList[i])
        composer.append(doc_temp)
    composer.save(finalPath)

def create_table11(path_subarea):
    path_excel = r"data\Cronograma Vissim.xlsx"
    df = pd.read_excel(path_excel, "Cronograma", header=0, usecols="A:E", skiprows=1)
    df = df.drop(columns=["Codigo TDR", "Entregable"])

    numsubarea = os.path.split(path_subarea)[1][-3:]
    no = int(numsubarea)

    code_by_subarea = df[df['Sub Area'] == no]['Codigo'].tolist()
    code_n_name = []
    for code in code_by_subarea:
        name_by_code = df[df['Codigo'] == code]['Intersección'].unique()[0]
        code_n_name.append([code, name_by_code])

    tablasPathList = []
    for code, name in code_n_name:
        doc = Document()
        table = doc.add_table(rows = 2, cols = 2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #Header:
        table.cell(0,0).text = name
        table.cell(0,0).merge(table.cell(0,1))

        for j in [0]:
            for i in range(len(table.columns)):
                cell_xml_element = table.rows[j].cells[i]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement('w:shd')
                shade_obj.set(qn('w:fill'),'B4C6E7')
                table_cell_properties.append(shade_obj)

        for row in table.rows:
            for i, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    try:
                        run = paragraph.runs[0]
                        run.font.name = 'Arial Narrow'
                        run.font.size = Pt(11)
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except IndexError:
                        pass
                    
        for selected_row in [0]:
            for cell in table.rows[selected_row].cells:
                for paragraph in cell.paragraphs:
                    try:
                        run = paragraph.runs[0]
                        run.font.bold = True
                    except IndexError:
                        pass
                    
        table.style = "Table Grid"

        tablaPath = os.path.join(path_subarea, "Tablas", f"tabla11_{code}.docx")
        doc.save(tablaPath)
        tablasPathList.append([code, tablaPath])

    wordsList = []
    for code, pathTable in tablasPathList:
        doc_template = DocxTemplate("./templates/template_tablas4.docx")
        texto = f"Fases semafóricas propuestas de la intersección {code}"
        newTable = doc_template.new_subdoc(pathTable)

        doc_template.render({
            "texto": texto,
            "tabla": newTable,
        })

        intermidiatePath = os.path.join(path_subarea, "Tablas", f"table11_{code}_REF.docx")
        doc_template.save(intermidiatePath)
        wordsList.append(intermidiatePath)

    finalPath = os.path.join(path_subarea, "Tablas", "table11.docx")
    filePathMaster = wordsList[0]
    filePathList = wordsList[1:]
    _combine_all_docx(filePathMaster, filePathList, finalPath)

    return finalPath