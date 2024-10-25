import os
from pathlib import Path
import re
from tables.tools.matrix import read_matrix
import pandas as pd
import xml.etree.ElementTree as ET
import numpy as np
from tqdm import tqdm
import time
#docx
from docxcompose.composer import Composer
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.enum.section import WD_SECTION

def _combine_all_docx(filePathMaster, filePathsList, finalPath) -> None:
    number_of_sections = len(filePathsList)
    master = Document(filePathMaster)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document(filePathsList[i])
        composer.append(doc_temp)

    composer.save(finalPath)

def _set_vertical_cell_direction(cell: _Cell, direction: str) -> None:
    #direction: tbRl -- Top to bottom, Right to left
    #direction: btLr -- Bottom to top, Left to right
    assert direction in ("tbRl", "btLr")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)
    tcPr.append(textDirection)

def table_creation(MATRIX, nameScenario, tipicidad, subareaPath):

    DESTINYS = MATRIX.columns.tolist()
    ORIGINS = MATRIX.index.tolist()

    doc = Document()
    table = doc.add_table(rows = 1+len(ORIGINS), cols = 1+len(DESTINYS))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr_cells = table.rows[0].cells
    for i, elem in enumerate(DESTINYS):
        #table.cell(0,i+1).text = str(elem)
        hdr_cells[i+1].text = str(elem)

    for i, elem in enumerate(ORIGINS):
        table.cell(i+1,0).text = str(elem)

    for i, row in enumerate(MATRIX.itertuples(index=False, name=None)):
        for j, elem in enumerate(row):
            cell = table.cell(i + 1, j + 1)
            cell.text = str(int(elem))

    # #Format
    # table.cell(2,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # _set_vertical_cell_direction(
    #     table.cell(2,0),
    #     "btLr",
    # )

    #Filling blank spaces:
    table.cell(0,0).text = "O\D"

    for selected_row in [0]:
        for cell in table.rows[selected_row].cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.bold = True

    for i in range(len(table.rows)):
        cell = table.cell(i,0)
        for paragraph in cell.paragraphs:
            run = paragraph.runs[0]
            run.font.bold = True
    
    for i in range(len(table.columns)):
        cell_xml_element = table.rows[0].cells[i]._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'),'B4C6E7')
        table_cell_properties.append(shade_obj)

    for i in range(len(table.rows)):
        cell = table.cell(i,0)
        cell_xml_element = cell._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'),'B4C6E7')
        table_cell_properties.append(shade_obj)

    for id in range(table.columns.__len__()):
        for cell in table.columns[id].cells:
            cell.width = Inches(0.25)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                try:
                    run = paragraph.runs[0]
                    run.font.name = 'Arial Narrow'
                    run.font.size = Pt(11)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except IndexError:
                    pass

    cell_exceptions = [(0,0),(0,1),(1,0),(1,1)]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if (i,j) in cell_exceptions: continue
            cell.style = 'Table Grid'

    if tipicidad == 'Tipico': tip = 'T'
    elif tipicidad == 'Atipico': tip = 'A'

    table.style = 'Table Grid'

    tablasFolder = os.path.join(subareaPath, 'Tablas')
    final_name = f"table14_{nameScenario}_{tip}.docx"
    finalPath = os.path.join(tablasFolder,final_name)
    doc.save(finalPath)

    return finalPath

def create_table14(path_subarea, maxTurn, maxTipicidad):
    subareaName = os.path.split(path_subarea)[-1]
    subareaID = subareaName[-3:]
    #subarea_id = path_parts[-1]
    actualFolder = Path(path_subarea) / "Actual"

    #Finding balanced scenario:
    balanceFolder = Path(path_subarea) / "Balanceado"
    for tipicidad in ["Tipico", "Atipico"]:
        folderList = os.listdir(balanceFolder / tipicidad)
        folderList = [file for file in folderList if not file.endswith('.ini')]
        for folder in folderList:
            pathFolder = balanceFolder / tipicidad / folder
            contentList = os.listdir(pathFolder)
            for content in contentList:
                if 'Reporte_GEH-R2' in content:
                    tipicidadBalanced = tipicidad
                    scenarioBalanced = folder
                    break

    if scenarioBalanced == "Manana": scenarioBalanced = "HPM"
    elif scenarioBalanced == "Tarde": scenarioBalanced = "HPT"
    elif scenarioBalanced == "Noche": scenarioBalanced = "HPN"
    
    #Unique matrix
    scenarioFolder = actualFolder / tipicidadBalanced / scenarioBalanced
    listContent = os.listdir(scenarioFolder)
    matrixExcel = [file for file in listContent if file.endswith(".xlsx") and 'Matriz-OD_' in file]

    assert len(matrixExcel) == 1, f"Hay más de una matriz o no hay en: {scenarioFolder}"
    matrixExcelPath = scenarioFolder / matrixExcel[0]
    ORIGINS, DESTINATIONS, MATRIX = read_matrix(matrixExcelPath)
    tablePath = table_creation(MATRIX, matrixExcel[0][:-5], tipicidadBalanced, path_subarea)

    docTemplate = DocxTemplate("./templates/template_tablas2.docx")
    texto = f"Matriz OD de la subarea {subareaID} {scenarioBalanced} día {tipicidadBalanced.lower()}"
    newTable = docTemplate.new_subdoc(tablePath)
    VARIABLES = {
        'texto': texto,
        'tabla': newTable,
    }
    docTemplate.render(VARIABLES)
    table14_path = os.path.join(path_subarea, "Tablas", "tabla14.docx")
    docTemplate.save(table14_path)

    selectedInformation = [ORIGINS, DESTINATIONS, MATRIX]

    dataframeMatrix = selectedInformation[2]

    max_value = dataframeMatrix.max().max()
    max_position = np.where(dataframeMatrix == max_value)
    rowMax, colMax = max_position[0][0], max_position[1][0]

    #New variables for the paragraph
    vissimFile = os.listdir(path_subarea)
    vissimFile = [file for file in vissimFile if file.endswith(".inpx")][0]
    vissimPath = os.path.join(path_subarea, vissimFile)
    tree = ET.parse(vissimPath)
    network_tag = tree.getroot()

    numberNodes = len(network_tag.findall("./nodes/node"))

    if numberNodes > 1:
        joinExplanation = "Las uniones de las intersecciones también son consideradas como centroides de OD. "
    else:
        joinExplanation = ""

    VARIABLES = {
        'numorig': str(len(selectedInformation[0])),
        'numdesti': str(len(selectedInformation[1])),
        'numorigmax': str(rowMax),
        'numdestimax': str(colMax),
        'joinExplanation': joinExplanation
    }

    #TODO: Anteriormente estaban las variables orig y desti, en los que se necesitaba
    #el nombre del link en donde se encuentra dicho origen o destino, la idea sería
    #conseguir ese dataframe y guardarlo para poder leerlo más adelante.

    table14_path = os.path.normpath(table14_path)

    return table14_path, VARIABLES