import os
import pandas as pd
from pathlib import Path

#docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_table1(path_subarea):
    path_excel = r"data\Cronograma Vissim.xlsx"
    df = pd.read_excel(path_excel, sheet_name='Cronograma', header=0, usecols="A:E", skiprows=1)
    df = df.drop(columns=["Codigo TDR", "Entregable"])

    doc = Document()
    numsubarea = os.path.split(path_subarea)[1][-3:]

    no = int(numsubarea)

    code_by_subarea = df[df['Sub Area'] == no]['Codigo'].tolist()
    name_by_subarea = df[df['Sub Area'] == no]['Intersección'].tolist()

    table = doc.add_table(rows = 1+len(code_by_subarea), cols = 3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #Headers:
    table.cell(0,0).text = "Código"
    table.cell(0,1).text = "Intersección"
    table.cell(0,2).text = "Subárea"

    for i in range(len(code_by_subarea)):
        table.cell(i+1, 0).text = code_by_subarea[i]
        table.cell(i+1, 1).text = name_by_subarea[i]
        table.cell(i+1, 2).text = numsubarea
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            run = paragraph.runs[0]
            run.font.bold = True

    for i in range(len(table.columns)):
        cell_xml_element = table.rows[0].cells[i]._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'),'B4C6E7')
        table_cell_properties.append(shade_obj)

    for id, x in zip([0,1,2],[0.5,3,0.5]):
        for cell in table.columns[id].cells:
            cell.width = Inches(x)
    table.style = 'Table Grid'

    final_path = Path(path_subarea) / "Tablas" / "table1.docx"

    if not os.path.exists(final_path.parent):
        os.makedirs(final_path.parent)

    doc.save(final_path)

    return final_path