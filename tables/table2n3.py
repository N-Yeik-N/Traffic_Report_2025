import os
import pandas as pd
from pathlib import Path
from tables.tools.peakfinder import peakhour_finder_and_volumes, compute_ph_system
from tables.tools.reading import *
import re
from dataclasses import dataclass

#docx
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _combine_all_docx(filePathMaster, filePathsList, finalPath) -> None:
    number_of_sections = len(filePathsList)
    master = Document(filePathMaster)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document(filePathsList[i])
        composer.append(doc_temp)

    composer.save(finalPath)

@dataclass
class Paragraph:
    tipicidad: str
    turno: str
    horaturno: str
    codigo: str
    maxvolveh: str

def create_table2_vehicular(
        path_subarea: str,
        code_by_subarea: list,
        peakHours: dict,
        listExcelData: dict,
        ):

    doc = Document()
    table = doc.add_table(rows = 5+len(code_by_subarea)*2+1, cols = 8)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #**** Headers *****#
    table.cell(0,0).text = "Código"
    table.cell(0,1).text = 'Intersección'

    table.cell(0,2).text = 'Turno Mañana'
    table.cell(0,2).merge(table.cell(0,3))

    table.cell(0,4).text = 'Turno Tarde'
    table.cell(0,4).merge(table.cell(0,5))

    table.cell(0,6).text = 'Turno Noche'
    table.cell(0,6).merge(table.cell(0,7))

    table.cell(0,0).merge(table.cell(1,0))
    table.cell(0,1).merge(table.cell(1,1))

    for i in range(2, 8, 2):
        table.cell(1, i).text = 'Hora Punta'
        table.cell(1, i+1).text = 'Vol.'

    #**** Separation ****#
    table.cell(2,0).text = "Día típico".upper()
    table.cell(2,0).merge(table.cell(2,7))

    #**** Separation of System Peak Hour Typical ****#
    hp_row1 = 1+len(code_by_subarea)+2
    table.cell(hp_row1,0).text = "Hora Punta"
    table.cell(hp_row1,0).merge(table.cell(hp_row1,1))

    peakHourSystem1 = str2hour(peakHours["Tipico"]["Morning"])
    table.cell(hp_row1,2).text = peakHourSystem1
    table.cell(hp_row1,2).merge(table.cell(hp_row1,3))

    peakHourSystem2 = str2hour(peakHours["Tipico"]["Evening"])
    table.cell(hp_row1,4).text = peakHourSystem2
    table.cell(hp_row1,4).merge(table.cell(hp_row1,5))

    peakHourSystem3 = str2hour(peakHours["Tipico"]["Night"])
    table.cell(hp_row1,6).text = peakHourSystem3
    table.cell(hp_row1,6).merge(table.cell(hp_row1,7))

    #**** Separation ****#
    table.cell(1+len(code_by_subarea)+3,0).text = "Día Atípico".upper()
    table.cell(1+len(code_by_subarea)+3,0).merge(table.cell(1+len(code_by_subarea)+3,7))

    #**** Separation of System Peak Hour Atypical ****#
    hp_row2 = 1+len(code_by_subarea)*2+4
    table.cell(hp_row2,0).text = "Hora Punta"
    table.cell(hp_row2,0).merge(table.cell(hp_row2,1))

    peakHourSystem1 = str2hour(peakHours["Atipico"]["Morning"])
    table.cell(hp_row2,2).text = peakHourSystem1
    table.cell(hp_row2,2).merge(table.cell(hp_row2,3))

    peakHourSystem2 = str2hour(peakHours["Atipico"]["Evening"])
    table.cell(hp_row2,4).text = peakHourSystem2
    table.cell(hp_row2,4).merge(table.cell(hp_row2,5))

    peakHourSystem3 = str2hour(peakHours["Atipico"]["Night"])
    table.cell(hp_row2,6).text = peakHourSystem3
    table.cell(hp_row2,6).merge(table.cell(hp_row2,7))

    start_tipico = 3
    start_atipico = 2+len(code_by_subarea)+3

    for i, node in enumerate(listExcelData["Tipico"]):
        table.cell(start_tipico+i,0).text = node.codigo
        table.cell(start_tipico+i,1).text = node.name

        ph1 = str2hour(node.id_morning)
        table.cell(start_tipico+i,2).text = ph1
        table.cell(start_tipico+i,3).text = str(node.vol_morning)

        ph2 = str2hour(node.id_evening)
        table.cell(start_tipico+i,4).text = ph2
        table.cell(start_tipico+i,5).text = str(node.vol_evening)

        ph3 = str2hour(node.id_night)
        table.cell(start_tipico+i,6).text = ph3
        table.cell(start_tipico+i,7).text = str(node.vol_night)
        
    for i, node in enumerate(listExcelData["Atipico"]):
        table.cell(start_atipico+i,0).text = node.codigo
        table.cell(start_atipico+i,1).text = node.name

        ph1 = str2hour(node.id_morning)
        table.cell(start_atipico+i,2).text = ph1
        table.cell(start_atipico+i,3).text = str(node.vol_morning)

        ph2 = str2hour(node.id_evening)
        table.cell(start_atipico+i,4).text = ph2
        table.cell(start_atipico+i,5).text = str(node.vol_evening)

        ph3 = str2hour(node.id_night)
        table.cell(start_atipico+i,6).text = ph3
        table.cell(start_atipico+i,7).text = str(node.vol_night)

    for selected_row in [0, 1, 2, hp_row1, hp_row1+1, hp_row2]: #hp_row2, 1+len(code_by_subarea)+2
        for cell in table.rows[selected_row].cells:
            for paragraph in cell.paragraphs:
                try:
                    run = paragraph.runs[0]
                except IndexError:
                    continue
                run.font.bold = True

    for selectedRow in [0,1]:
        for i in range(len(table.columns)):
            cell_xml_element = table.rows[selectedRow].cells[i]._tc
            table_cell_properties = cell_xml_element.get_or_add_tcPr()
            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'),'B4C6E7')
            table_cell_properties.append(shade_obj)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                try:
                    run = paragraph.runs[0]
                except IndexError:
                    continue
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for id, x in zip(
        [0,1,2,3,4,5,6,7],
        [1.48,4.06,2.66,1,2.66,1,2.66,1]):
        for cell in table.columns[id].cells:
            cell.width = Cm(x)

    table.style = 'Table Grid'

    finalPath2_vehicular = Path(path_subarea) / "Tablas" / "table2_vehicular.docx"

    doc.save(finalPath2_vehicular)

    return finalPath2_vehicular #-----> dconteot, dconteoa

def create_table2_peatonal(
        path_subarea: str,
        code_by_subarea: list,
        peakHours: dict,
        listPeakHoursPed: list,
) -> str:
    
    doc = Document()
    table = doc.add_table(rows = 5+len(code_by_subarea)*2+1, cols = 8)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #**** Headers *****#
    table.cell(0,0).text = "Código"
    table.cell(0,1).text = 'Intersección'

    table.cell(0,2).text = 'Turno Mañana'
    table.cell(0,2).merge(table.cell(0,3))

    table.cell(0,4).text = 'Turno Tarde'
    table.cell(0,4).merge(table.cell(0,5))

    table.cell(0,6).text = 'Turno Noche'
    table.cell(0,6).merge(table.cell(0,7))

    table.cell(0,0).merge(table.cell(1,0))
    table.cell(0,1).merge(table.cell(1,1))

    for i in range(2, 8, 2):
        table.cell(1, i).text = 'Hora Punta'
        table.cell(1, i+1).text = 'Vol.'

    #**** Separation ****#
    table.cell(2,0).text = "Día típico".upper()
    table.cell(2,0).merge(table.cell(2,7))

    #**** Separation of System Peak Hour Typical ****#
    hp_row1 = 1+len(code_by_subarea)+2
    table.cell(hp_row1,0).text = "Hora Punta"
    table.cell(hp_row1,0).merge(table.cell(hp_row1,1))

    peakHourSystem1 = str2hour(peakHours["Tipico"]["Morning"])
    table.cell(hp_row1,2).text = peakHourSystem1
    table.cell(hp_row1,2).merge(table.cell(hp_row1,3))

    peakHourSystem2 = str2hour(peakHours["Tipico"]["Evening"])
    table.cell(hp_row1,4).text = peakHourSystem2
    table.cell(hp_row1,4).merge(table.cell(hp_row1,5))

    peakHourSystem3 = str2hour(peakHours["Tipico"]["Night"])
    table.cell(hp_row1,6).text = peakHourSystem3
    table.cell(hp_row1,6).merge(table.cell(hp_row1,7))

    #**** Separation ****#
    table.cell(1+len(code_by_subarea)+3,0).text = "Día Atípico".upper()
    table.cell(1+len(code_by_subarea)+3,0).merge(table.cell(1+len(code_by_subarea)+3,7))

    #**** Separation of System Peak Hour Atypical ****#
    hp_row2 = 1+len(code_by_subarea)*2+4
    table.cell(hp_row2,0).text = "Hora Punta"
    table.cell(hp_row2,0).merge(table.cell(hp_row2,1))

    peakHourSystem1 = str2hour(peakHours["Atipico"]["Morning"])
    table.cell(hp_row2,2).text = peakHourSystem1
    table.cell(hp_row2,2).merge(table.cell(hp_row2,3))

    peakHourSystem2 = str2hour(peakHours["Atipico"]["Evening"])
    table.cell(hp_row2,4).text = peakHourSystem2
    table.cell(hp_row2,4).merge(table.cell(hp_row2,5))

    peakHourSystem3 = str2hour(peakHours["Atipico"]["Night"])
    table.cell(hp_row2,6).text = peakHourSystem3
    table.cell(hp_row2,6).merge(table.cell(hp_row2,7))

    start_tipico = 3
    start_atipico = 2+len(code_by_subarea)+3

    #**** Individual peak hours ****#
    for i, node in enumerate(listPeakHoursPed["Tipico"]):
        table.cell(start_tipico+i,0).text = node.code
        table.cell(start_tipico+i,1).text = node.name

        ph1 = str2hour(node.idMorning)
        table.cell(start_tipico+i,2).text = ph1
        table.cell(start_tipico+i,3).text = str(node.morningVolume)

        ph2 = str2hour(node.idEvening)
        table.cell(start_tipico+i,4).text = ph2
        table.cell(start_tipico+i,5).text = str(node.eveningVolume)

        ph3 = str2hour(node.idNight)
        table.cell(start_tipico+i,6).text = ph3
        table.cell(start_tipico+i,7).text = str(node.nightVolume)
        
    for i, node in enumerate(listPeakHoursPed["Atipico"]):
        table.cell(start_atipico+i,0).text = node.code
        table.cell(start_atipico+i,1).text = node.name

        ph1 = str2hour(node.idMorning)
        table.cell(start_atipico+i,2).text = ph1
        table.cell(start_atipico+i,3).text = str(node.morningVolume)

        ph2 = str2hour(node.idEvening)
        table.cell(start_atipico+i,4).text = ph2
        table.cell(start_atipico+i,5).text = str(node.eveningVolume)

        ph3 = str2hour(node.idNight)
        table.cell(start_atipico+i,6).text = ph3
        table.cell(start_atipico+i,7).text = str(node.nightVolume)

    for selected_row in [0, 1, 2, hp_row1, hp_row1+1, hp_row2]: #hp_row2, 1+len(code_by_subarea)+2
        for cell in table.rows[selected_row].cells:
            for paragraph in cell.paragraphs:
                try:
                    run = paragraph.runs[0]
                except IndexError:
                    continue
                run.font.bold = True

    for selectedRow in [0,1]:
        for i in range(len(table.columns)):
            cell_xml_element = table.rows[selectedRow].cells[i]._tc
            table_cell_properties = cell_xml_element.get_or_add_tcPr()
            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'),'B4C6E7')
            table_cell_properties.append(shade_obj)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                try:
                    run = paragraph.runs[0]
                except IndexError:
                    continue
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for id, x in zip(
        [0,1,2,3,4,5,6,7],
        [1.48,4.06,2.66,1,2.66,1,2.66,1]):
        for cell in table.columns[id].cells:
            cell.width = Cm(x)

    table.style = 'Table Grid'

    finalPath2_peatonal = Path(path_subarea) / "Tablas" / "table2_peatonal.docx"

    doc.save(finalPath2_peatonal)

    return finalPath2_peatonal

def create_table3(path_subarea, dconteot, dconteoa, codeBySubarea):
    doc = Document()
    table = doc.add_table(rows=7 ,cols=5)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table.cell(0,0).text = "Intersección"
    table.cell(0,1).text = "Día"
    table.cell(0,2).text = "Tipicidad"
    table.cell(0,3).text = "Turno"
    table.cell(0,4).text = "Horario"

    #codinterseccion:
    set_codes = ""
    for i, code in enumerate(codeBySubarea):
        if i == len(codeBySubarea) - 1:
            set_codes += code
        else:
            set_codes += code+'\n' 
    
    table.cell(1,0).text = set_codes
    table.cell(1,0).merge(table.cell(6,0))

    for i in range(3):
        table.cell(1+i,2).text = "Típico"
        table.cell(4+i,2).text = "Atípico"

    for i, valor in enumerate(["Mañana", "Tarde", "Noche"]):
        table.cell(1+i,3).text = valor
        table.cell(4+i,3).text = valor

    for i, valor in enumerate(["06:30 - 09:30",
                "12:00 - 15:00",
                "17:30 - 20:30",
                "06:30 - 09:30",
                "12:00 - 15:00",
                "17:30 - 20:30",]):
        table.cell(1+i,4).text = valor

    for i in range(3):
        table.cell(1+i,1).text = dconteot
        table.cell(4+i,1).text = dconteoa

    for selected_row in [0]:
        for cell in table.rows[selected_row].cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.bold = True

    for i in range(len(table.columns)):
        cell_xml_element = table.rows[0].cells[i]._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'),'B4C6E7')
        table_cell_properties.append(shade_obj)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for id, x in zip([0,1,2,3,4],[0.5,1,1,0.8,1.2]):
        for cell in table.columns[id].cells:
            cell.width = Inches(x)

    table.style = 'Table Grid'

    final_path_3 = Path(path_subarea) / "Tablas" / "table3.docx"

    doc.save(final_path_3)

    return final_path_3

def create_tables2n3(pathSubarea: str):
    #######################
    # Vehicle Information #
    #######################
    
    #codes by list
    codeBySubarea = code_by_subarea(pathSubarea)

    #List of excels
    pathParts = pathSubarea.split("/")
    subareaID = pathParts[-1]
    proyectFolder = '/'.join(pathParts[:-2])
    fieldData = Path(proyectFolder) / "7. Informacion de Campo" / subareaID / "Vehicular"

    excelByTipicidad = {}
    for tipicidad in ["Tipico", "Atipico"]:
        typicalPath = fieldData / tipicidad
        listExcels = os.listdir(typicalPath)
        listExcels = [str(typicalPath / file) for file in listExcels if file.endswith(".xlsm") and not file.startswith("~")]
        excelByTipicidad[tipicidad] = listExcels

    listExcelData = {
        "Tipico": [],
        "Atipico": [],
    }

    dayData = {
        "Tipico": [],
        "Atipico": [],
    }

    totalSumCounts = [0 for _ in range(96)]
    totalSumDict = {
        "Tipico": totalSumCounts,
        "Atipico": totalSumCounts,
    }

    for tipicidad, listExcels in excelByTipicidad.items():
        for excelPath in listExcels:
            excelHourInfo, sumQuarter = peakhour_finder_and_volumes(excelPath)
            totalSumDict[tipicidad] = [x+y for x,y in zip(totalSumDict[tipicidad], sumQuarter)]
            listExcelData[tipicidad].append(excelHourInfo)
            dayData[tipicidad].append(excelHourInfo.fecha)

    peakHours = {
        "Tipico": {},
        "Atipico": {},
    }

    volSums = {
        "Tipico": {},
        "Atipico": {},
    }

    for tipicidad, listHourData in listExcelData.items():
        #NOTE: Level: Typical o Atypical
        MORNING = []
        EVENING = []
        NIGHT   = []
        for excelHourData in listHourData:
            MORNING.append((excelHourData.id_morning, excelHourData.vol_morning))
            EVENING.append((excelHourData.id_evening, excelHourData.vol_evening))
            NIGHT.append((excelHourData.id_night, excelHourData.vol_night))

        hourSystem1 = compute_ph_system(MORNING)
        hourSystem2 = compute_ph_system(EVENING)
        hourSystem3 = compute_ph_system(NIGHT)

        peakHours[tipicidad].update({
            "Morning": hourSystem1,
            "Evening": hourSystem2,                    
            "Night": hourSystem3,
        })

        #Total volumes:
        volSum1 = int(sum(totalSumDict[tipicidad][hourSystem1:hourSystem1+4]))
        volSum2 = int(sum(totalSumDict[tipicidad][hourSystem2:hourSystem2+4]))
        volSum3 = int(sum(totalSumDict[tipicidad][hourSystem3:hourSystem3+4]))

        volSums[tipicidad].update({
            "Morning": volSum1,
            "Evening": volSum2,
            "Night": volSum3,
        })

    #Preparing paragraphs:
    paragraph_dict = {
        "Tipico": {
            "Mañana": None,
            "Tarde": None,
            "Noche": None,
        },
        "Atipico": {
            "Mañana": None,
            "Tarde": None,
            "Noche": None,
        }
    }

    if len(codeBySubarea) == 1:
        codigoTxt = codeBySubarea[0]
    else:
        codigoTxt = ', '.join(codeBySubarea[:-1]) + ' y ' + codeBySubarea[-1]

    for tipicidad in ["Tipico", "Atipico"]:
        for hour, hourEnglish in zip(["Mañana", "Tarde", "Noche"], ["Morning", "Evening", "Night"]):
            paragraph_dict[tipicidad][hour] = Paragraph(
                tipicidad = tipicidad,
                turno = hour,
                horaturno = str2hour(peakHours[tipicidad][hourEnglish]),
                codigo = codigoTxt,
                maxvolveh = str(volSums[tipicidad][hourEnglish]),
            )

    countParagraphs = 0
    listParagraphsPaths = []
    for tipicidad in ["Tipico", "Atipico"]:
        for hour in ["Mañana", "Tarde", "Noche"]:
            docTemplate = DocxTemplate("./templates/template_lista3.docx")
            dataParagraph = paragraph_dict[tipicidad][hour]
            if tipicidad == "Tipico":
                tipicidadtxt = "típico"
            else:
                tipicidadtxt = "atípico"
            docTemplate.render({
                "tipicidad_ph": tipicidadtxt,
                "turno_ph": dataParagraph.turno.lower(),
                "horaturno": dataParagraph.horaturno,
                "codintersection_ph": dataParagraph.codigo,
                "maxvolveh": dataParagraph.maxvolveh
            })
            finalPath = os.path.join(
                pathSubarea, "Tablas", f"paragraph_ph_{countParagraphs}.docx"
            )
            docTemplate.save(finalPath)
            listParagraphsPaths.append(finalPath)
            countParagraphs += 1

    filePathMaster = listParagraphsPaths[0]
    filePathList = listParagraphsPaths[1:]
    paragraph_ph_path = os.path.join(pathSubarea, "Tablas", "paragraph_ph.docx")
    _combine_all_docx(filePathMaster, filePathList, paragraph_ph_path)

    #######################
    # Creating paragraphs #
    #######################

    #Creating .txt with peakhours:
    pathSubarea = Path(pathSubarea)
    contentTipico = f"Morning:\t{peakHours['Tipico']['Morning']/4}\nEvening:\t{peakHours['Tipico']['Evening']/4}\nNight:\t{peakHours['Tipico']['Night']/4}"
    peakhourstip_path = pathSubarea / "Tablas" / "PeakHoursTipico.txt"
    with open(peakhourstip_path, "w") as file:
        file.write(contentTipico)

    contentTipico = f"Morning:\t{peakHours['Atipico']['Morning']/4}\nEvening:\t{peakHours['Atipico']['Evening']/4}\nNight:\t{peakHours['Atipico']['Night']/4}"
    peakhoursati_path = pathSubarea / "Tablas" / "PeakHoursAtipico.txt"
    with open(peakhoursati_path, "w") as file:
        file.write(contentTipico)

    ##########################
    # Pedestrian Information #
    ##########################

    fieldData = Path(proyectFolder) / "7. Informacion de Campo" / subareaID / "Peatonal"
    excelByTipicidad = {}
    for tipicidad in ["Tipico", "Atipico"]:
        typicalPath = fieldData / tipicidad
        listExcels = os.listdir(typicalPath)
        listExcels = [str(typicalPath / file) for file in listExcels if file.endswith(".xlsm") and not file.startswith("~")]
        excelByTipicidad[tipicidad] = listExcels

    listPedData = {
        "Tipico": [],
        "Atipico": [],
    }

    pattern = r"([A-Z]+-[0-9]+)"


    volPedDict = {
        "Tipico": totalSumCounts,
        "Atipico": totalSumCounts,
    }

    for tipicidad, listExcels in excelByTipicidad.items():
        for excelPath in listExcels:
            excelName = os.path.split(excelPath)[1][:-5]
            excelCodigo = re.search(pattern, excelName).group(1)

            columVolumes, volumesPed = read_ped_excel_and_volumes(excelPath)
            data = PedestrianVolumes(
                codigo = excelCodigo,
                volTotal = columVolumes,
            )

            volPedDict[tipicidad] = [x+y for x, y in zip(volPedDict[tipicidad], volumesPed)]
            listPedData[tipicidad].append(data)

    listPeakHoursPed = { #TODO: Check if this is correct
        "Tipico": [],
        "Atipico": [],
    }

    for tipicidad, dataPEDS in listPedData.items():
        listExcelVeh = listExcelData[tipicidad]
        for dataPED in dataPEDS:
            for dataVEH in listExcelVeh:
                if dataPED.codigo == dataVEH.codigo:
                    data = PedestrianInfo(
                        code = dataVEH.codigo,
                        name = dataVEH.name,
                        idMorning = dataVEH.id_morning,
                        idEvening = dataVEH.id_evening,
                        idNight = dataVEH.id_night,
                        morningVolume = dataPED.volTotal[dataVEH.id_morning+3],
                        eveningVolume = dataPED.volTotal[dataVEH.id_evening+3],
                        nightVolume = dataPED.volTotal[dataVEH.id_night+3],
                    )
                    listPeakHoursPed[tipicidad].append(data)
                    break

    ######################################
    # Creating paragraphs for pedestrian #
    ######################################

    paragraph_dict_ped = {
        "Tipico": {},
        "Atipico": {},
    }

    for tipicidad in ["Tipico", "Atipico"]:
        ph1 = peakHours[tipicidad]["Morning"]
        ph2 = peakHours[tipicidad]["Evening"]
        ph3 = peakHours[tipicidad]["Night"]
        volPedSum1 = int(sum(volPedDict[tipicidad][ph1:ph1+4]))
        volPedSum2 = int(sum(volPedDict[tipicidad][ph2:ph2+4]))
        volPedSum3 = int(sum(volPedDict[tipicidad][ph3:ph3+4]))

        paragraph_dict_ped[tipicidad].update({
            "Mañana": volPedSum1,
            "Tarde": volPedSum2,
            "Noche": volPedSum3,
        })

    countParagraphs = 0
    listParagraphsPeds = []
    for tipicidad in ["Tipico", "Atipico"]:
        for hour in ["Mañana", "Tarde", "Noche"]:
            docTemplate = DocxTemplate("./templates/template_lista4.docx")
            maxVolPed = paragraph_dict_ped[tipicidad][hour]
            dataParagraph = paragraph_dict[tipicidad][hour]
            if tipicidad == "Tipico":
                tipicidadtxt = "típico"
            else:
                tipicidadtxt = "atípico"

            docTemplate.render({
                "tipicidad_ph": tipicidadtxt,
                "turno_ph": hour.lower(),
                "horaturno": dataParagraph.horaturno,
                "codintersection_ph": dataParagraph.codigo,
                "maxvolped": str(maxVolPed),
            })

            finalPathPed = os.path.join(
                pathSubarea, "Tablas", f"paragraph_ph_ped_{countParagraphs}.docx"
            )

            docTemplate.save(finalPathPed)
            listParagraphsPeds.append(finalPathPed)
            countParagraphs += 1

    filePathMaster = listParagraphsPeds[0]
    filePathList = listParagraphsPeds[1:]
    paragraph_ph_ped = os.path.join(pathSubarea, "Tablas", "paragraph_ph_ped.docx")
    _combine_all_docx(filePathMaster, filePathList, paragraph_ph_ped)

    ###############################
    # Creating table 2: vehicular #
    ###############################

    try:
        finalPath2_vehicular = create_table2_vehicular(
            pathSubarea,
            codeBySubarea,
            peakHours,
            listExcelData,
        )
    except Exception as inst:
        print("Error - HPs Vehicular: ", inst)
        finalPath2_vehicular = None

    ##############################
    # Creating table 2: peatonal #
    ##############################

    try:
        finalPath2_peatonal = create_table2_peatonal(
            pathSubarea,
            codeBySubarea,
            peakHours,
            listPeakHoursPed,
        )
    except Exception as inst:
        print("Error - HPs Peatonal: ", inst)
        finalPath2_peatonal = None

    #################
    # Día de conteo #
    #################

    typicalDay = list(set(dayData["Tipico"]))[0]
    atypicalDay = list(set(dayData["Atipico"]))[0]

    typicalDay = typicalDay.strftime("%d/%m/%Y")
    atypicalDay = atypicalDay.strftime("%d/%m/%Y")

    ####################
    # Creating table 3 #
    ####################

    try:
        finalPath3 = create_table3(
            pathSubarea,
            typicalDay,
            atypicalDay,
            codeBySubarea,
        )
    except Exception as inst:
        print("Error - Fechas de conteo: ", inst)
        finalPath3 = None

    return finalPath2_vehicular, finalPath2_peatonal, finalPath3, typicalDay, atypicalDay, paragraph_ph_path, paragraph_ph_ped