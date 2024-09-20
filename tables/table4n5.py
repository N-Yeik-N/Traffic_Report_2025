import os
from tables.tools.tale import tale_by_excel
#from tools.tale import tale_by_excel
from pathlib import Path

#docx
from docxcompose.composer import Composer
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _combine_all_docx(filePathMaster, filePathsList, finalPath) -> None:
    number_of_sections = len(filePathsList)
    master = Document(filePathMaster)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document(filePathsList[i])
        composer.append(doc_temp)

    composer.save(finalPath)

def append_document_content(source_doc, target_doc) -> None:
    for element in source_doc.element.body:
        target_doc.element.body.append(element)

def create_table4n5(path_subarea):
    path_subarea = Path(path_subarea)
    subarea_id = path_subarea.name
    proyect_folder = path_subarea.parents[1]

    field_data = os.path.join(
        proyect_folder,
        "7. Informacion de Campo",
        subarea_id,
        "Longitud de Cola",
    )

    excels_by_tipicidad = {}
    for tipicidad in ["Tipico", "Atipico"]:
        tipicidad_folder = os.path.join(field_data, tipicidad)
        list_excels = os.listdir(tipicidad_folder)
        list_excels = [os.path.join(tipicidad_folder, file) for file in list_excels
                       if file.endswith(".xlsx") and not file.startswith("~")]
        excels_by_tipicidad[tipicidad] = list_excels

    data_by_tipicidad = {}
    list_codes = []
    for tipicidad, excelList in excels_by_tipicidad.items():
        dataList = []
        for excel in excelList:
            codigo, date, dict_info, name = tale_by_excel(excel)
            list_codes.append(codigo)
            dataList.append([codigo, date, dict_info])
        data_by_tipicidad[tipicidad] = dataList

    ###############
    # Queues list #
    ###############

    generalTextSingular = "En el turno {TURN}, el sentido {SENT} presenta una longitud de cola máxima de {LONGCOLAMAX}m, siendo el sentido más crítico ({NUMAUTOS} autos en cola aproximadamente)."
    generalTextPlural = "En el turno {TURN}, los sentidos {SENT} presentan una longitud de cola máxima de {LONGCOLAMAX}m, siendo los sentidos más críticos ({NUMAUTOS} autos en cola aproximadamente)."

    if not os.path.exists(os.path.join(
        path_subarea, "Tablas", "Queues"
    )):
        os.makedirs(os.path.join(
            path_subarea, "Tablas", "Queues"
        ))

    queueWordLists = []
    for codigo, _, df, name in dataList:
        queueDict = {}
        max_values = df.groupby('Turn')['Max'].transform('max')
        dfResult = df[df['Max'] == max_values]
        for turn in ["Mañana", "Tarde", "Noche"]:
            dfTurn = dfResult[dfResult['Turn'] == turn].reset_index(drop=True)
            if dfTurn.shape[0] == 0:
                queueDict[turn] = f"No hay datos para el turno {turn.lower()}."
            elif dfTurn.shape[0] == 1:
                SENT = dfTurn.iloc[0]['Direction']
                LONGCOLAMAX = int(dfTurn.iloc[0]['Max'])
                NUMAUTOS = LONGCOLAMAX//5.5
                variables = {
                    "TURN": turn.lower(),
                    "SENT": SENT.lower(),
                    "LONGCOLAMAX": LONGCOLAMAX,
                    "NUMAUTOS": int(NUMAUTOS),
                }
                resultText = generalTextSingular.format(**variables)

            elif dfTurn.shape[0] >= 2:
                listSENT = [dfTurn.iloc[i]['Direction'] for i in range(dfTurn.shape[0])]
                if len(listSENT) > 2:
                    SENT = ', '.join(listSENT[:-1]) + ' y ' + listSENT[-1]
                elif len(listSENT) == 2:
                    SENT = ' y '.join(listSENT)
                LONGCOLAMAX = int(dfTurn.iloc[0]['Max'])
                NUMAUTOS = LONGCOLAMAX//5.5
                variables = {
                    "TURN": turn.lower(),
                    "SENT": SENT.lower(),
                    "LONGCOLAMAX": LONGCOLAMAX,
                    "NUMAUTOS": int(NUMAUTOS),
                } 
                resultText = generalTextPlural.format(**variables)

            queueDict[turn] = resultText

        doc = DocxTemplate("./templates/template_queue.docx")
        VARIABLES = {
            "codinterseccion": codigo,
            "nominterseccion": name,
            "morning": queueDict["mañana"],
            "afternoon": queueDict["tarde"],
            "night": queueDict["noche"],
        }
        doc.render(VARIABLES)
        queuePath = os.path.join(
            path_subarea, "Tablas", "Queues", f"{codigo}.docx"
        )
        doc.save(queuePath)
        queueWordLists.append(queuePath)
    
    finalPathQueue = os.path.join(path_subarea, "Tablas", "queues_list.docx")
    if len(queueWordLists) == 1:
        finalPathQueue = queueWordLists[0]
    elif queueWordLists > 1:
        filePathMaster = queueWordLists[0]
        filePathsList = queueWordLists[1:]
        _combine_all_docx(filePathMaster, filePathsList, finalPathQueue)
    else:
        finalPathQueue = None

    list_codes = list(set(list_codes))
    
    tipico_date = []
    atipico_date = []
    for tipicidad, dataList in data_by_tipicidad.items():
        for data in dataList:
            if tipicidad == "Tipico":
                tipico_date.append(data[1].strftime('%d/%m/%Y'))
            elif tipicidad == "Atipico":
                atipico_date.append(data[1].strftime('%d/%m/%Y'))

    tipico_date = list(set(tipico_date))
    atipico_date = list(set(atipico_date))

    ####################
    # Creating table 4 #
    ####################

    doc = Document()
    table = doc.add_table(rows=7, cols=5)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, texto in enumerate([
        "Intersección","Día","Tipicidad","Turno","Horario"]):
        table.cell(0,i).text = texto

    texto_codes = ""
    for i, code in enumerate(list_codes):
        if i == len(list_codes)-1:
            texto_codes += code
        else:
            texto_codes += code + "\n"

    table.cell(1,0).text = texto_codes
    table.cell(1,0).merge(table.cell(6,0))

    for i in range(3):
        table.cell(1+i,1).text = tipico_date[0]
        table.cell(4+i,1).text = atipico_date[0]

    for i in range(3):
        table.cell(1+i,2).text = "Tipico"
        table.cell(4+i,2).text = "Atipico"

    for i in range(2):
        table.cell(1+3*i,3).text = "Mañana"
        table.cell(2+3*i,3).text = "Tarde"
        table.cell(3+3*i,3).text = "Noche"

    for i, texto in enumerate([
        "06:30 - 09:30",
        "12:00 - 15:00",
        "17:30 - 20:30",
        "06:30 - 09:30",
        "12:00 - 15:00",
        "17:30 - 20:30",
    ]):
        table.cell(1+i,4).text = texto

    for selected_row in [0]:
        for cell in table.rows[selected_row].cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.bold = True

    for i in range(len(table.columns)):
        cell_xml_element = table.rows[0].cells[i]._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'),'B4C6E7')
        table_cell_properties.append(shade_obj)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for id, x in zip([0,1,2,3],[0.5,0.5,0.5,0.5]):
        for cell in table.columns[id].cells:
            cell.width = Inches(x)

    table.style = "Table Grid"
    table4_path = Path(path_subarea) / "Tablas" / "table4.docx"
    doc.save(table4_path)

    ####################
    # Creating table 5 #
    ####################

    infoConclusions = {
        "Tipico": {},
        "Atipico": {},
    }

    count = 1
    list_REF = []
    for tipicidad, dataList in data_by_tipicidad.items():
        for data in dataList:
            df = data[2]

            #Obtaining maximum queue:
            maxRowIndex = df['Max'].idxmax()
            maxRow = df.loc[maxRowIndex]

            infoConclusions[tipicidad][data[0]] = {
                "max_direction": maxRow['Direction'],
                "max_access": maxRow['Access'],
                "max_queue": maxRow['Max'],
            }

            if tipicidad == "Tipico": #NOTE: Why only Tipico?

                doc = Document()
                
                table = doc.add_table(rows=1, cols=6)
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for i, texto in enumerate([
                    "Turno","Sentido","Acceso","Longitud de Cola Máxima","Longitud de Cola Promedio","Desviación Estándar"
                ]):
                    table.cell(0,i).text = texto

                for i in range(len(df)):
                    new_row = table.add_row().cells
                    for j in range(len(df.columns)):
                        if j == 0: continue
                        try:
                            valor = f"{df.iloc[i,j]:.2f}"
                        except:
                            valor = str(df.iloc[i,j])
                        new_row[j].text = valor

                list_elem  = df['Turn'].value_counts()
                numbers_of_row = []
                for _, value in list_elem.items():
                    numbers_of_row.append(value)

                table.cell(1,0).text = "Mañana"
                table.cell(1,0).merge(table.cell(numbers_of_row[0],0))
                table.cell(numbers_of_row[0]+1,0).text = "Tarde"
                table.cell(numbers_of_row[0]+1,0).merge(table.cell(sum(numbers_of_row[:2]),0))
                table.cell(sum(numbers_of_row[:2])+1,0).text = "Noche"
                table.cell(sum(numbers_of_row[:2])+1,0).merge(table.cell(sum(numbers_of_row),0))

                #Encabezados en negrita
                for selected_row in [0]:
                    for cell in table.rows[selected_row].cells:
                        for paragraph in cell.paragraphs:
                            run = paragraph.runs[0]
                            run.font.bold = True

                #Colores de fondo
                for i in range(len(table.columns)):
                    cell_xml_element = table.rows[0].cells[i]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement('w:shd')
                    shade_obj.set(qn('w:fill'),'B4C6E7')
                    table_cell_properties.append(shade_obj)

                #Centrado
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            run = paragraph.runs[0]
                            run.font.name = 'Arial Narrow'
                            run.font.size = Pt(11)
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                #Ancho de columnas
                for id, x in zip([0,1,2,3,4,5],[0.5,0.5,2,1,1,1]):
                    for cell in table.columns[id].cells:
                        cell.width = Inches(x)
                
                table.style = 'Table Grid'
                table5_path = Path(path_subarea) / "Tablas" / f"table5_{count}.docx"
                doc.save(table5_path)
                
                doc_template = DocxTemplate("./templates/template_tablas.docx")
                new_table = doc_template.new_subdoc(table5_path)
                texto = f"Resumen de la longitud de cola de la intersección {data[0]} día típico"
                VARIABLES = {
                    'texto': texto,
                    'tabla': new_table,
                }
                doc_template.render(VARIABLES)
                ref_path = Path(path_subarea) / "Tablas" / f"table5_{count}_REF.docx"
                doc_template.save(ref_path)
                list_REF.append(ref_path)
                count += 1
    
    doc_target = Document(list_REF[0])
    for i in range(len(list_REF)):
        if i == 0: continue
        doc_source = Document(list_REF[i])
        append_document_content(doc_source, doc_target)
        table5_path_aux = Path(path_subarea) / "Tablas" / f"table5.docx"
        doc_target.save(table5_path_aux)
        doc_target = Document(table5_path_aux)

    table5_path = Path(path_subarea) / "Tablas" / "table5.docx"
    doc_target.save(table5_path)

    return table4_path, table5_path, finalPathQueue