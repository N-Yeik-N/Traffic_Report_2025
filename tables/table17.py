import os
from openpyxl import load_workbook
#docx
from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_table17(subareaPath):
    balancePath = os.path.join(subareaPath, "Balanceado")
    contentBalance = os.listdir(balancePath)
    contentBalance = [file for file in contentBalance if not file.endswith(".ini")]
    for tipicidad in contentBalance:
        tipicidadFolder = os.path.join(balancePath, tipicidad)
        contentTipicidad = os.listdir(tipicidadFolder)
        contentTipicidad = [file for file in contentTipicidad if not file.endswith(".ini")]
        for turnFolder in contentTipicidad:
            contentTurn = os.listdir(os.path.join(tipicidadFolder, turnFolder))
            for content in contentTurn:
                if content == "Reporte_GEH-R2.xlsm":
                    pathGEH = os.path.join(tipicidadFolder, turnFolder, content)
                    break
    
    assert pathGEH, "No se encontro el archivo GEH-R2.xlsm"

    wb = load_workbook(pathGEH, read_only=True, data_only=True)
    ws = wb.active

    typesNames = [row[0].value for row in ws[slice("H8","H27")] if row[0].value != None and row[0].value != 'n']
    typesNumbers = len(typesNames)

    listSlices = [
        slice("EJ8", "EJ27"),
        slice("EK8", "EK27"),
        slice("EL8", "EL27"),
        slice("EM8", "EM27"),
        slice("EN8", "EN27"),
    ]

    typesNames = typesNames[:typesNumbers]
    gehValues = [str(round(row[0].value,1)) for row in ws[listSlices[0]] if type(row[0].value) == float or type(row[0].value) == int][:typesNumbers]
    try:
        criterion1_1 =  ["{:.0%}".format(row[0].value) for row in ws[listSlices[1]]][:typesNumbers]
        criterion1_2 =  ["{:.0%}".format(row[0].value) for row in ws[listSlices[2]]][:typesNumbers]
        criterion1_3 =  ["{:.0%}".format(row[0].value) for row in ws[listSlices[3]]][:typesNumbers]
    except TypeError as e:
        print("Tabla 17\tERROR\tResultados del GEH-R2")
    try:
        criterion2_1 =  ["{:.0%}".format(row[0].value) for row in ws[listSlices[4]]][:typesNumbers]
    except ValueError as e:
        print("Extiende las fórmulas en el excel de GEH-R2, la zona de resultados. No debe haber !DIV=0")
        raise e

    wb.close()

    doc = Document()
    table = doc.add_table(rows = 1+typesNumbers, cols = 7)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #Headers:
    for i, texto in enumerate(["Tipología", "GEH", "Criterio 1"]):
        table.cell(0,i).text = texto

    for i, texto in enumerate(["Criterio 2", "Estado"]):
        table.cell(0,i+5).text = texto

    table.cell(0,2).merge(table.cell(0,4)) #Merge Criteria 1

    for i, texto in enumerate(typesNames):
        table.cell(i+1,0).text = texto

    for i, geh in enumerate(gehValues):
        table.cell(i+1,1).text = f"{geh:.2f}"

    for i, valor1_1 in enumerate(criterion1_1):
        table.cell(i+1,2).text = valor1_1
    
    for i, valor1_2 in enumerate(criterion1_2):
        table.cell(i+1,3).text = valor1_2

    for i, valor1_3 in enumerate(criterion1_3):
        table.cell(i+1,4).text = valor1_3

    for i, valor2_1 in enumerate(criterion2_1):
        table.cell(i+1,5).text = valor2_1

    for i in range(typesNumbers):
        table.cell(i+1,6).text = "Cumple"

    for selected_row in [0]:
            for cell in table.rows[selected_row].cells:
                for paragraph in cell.paragraphs:
                    try:
                        run = paragraph.runs[0]
                        run.font.bold = True
                    except IndexError:
                        pass

    for j in [0]:
            for i in range(len(table.columns)):
                cell_xml_element = table.rows[j].cells[i]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement('w:shd')
                shade_obj.set(qn('w:fill'),'B4C6E7')
                table_cell_properties.append(shade_obj)

    for row in table.rows:
            for i, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    try:
                        run = paragraph.runs[0]
                        run.font.name = 'Arial Narrow'
                        run.font.size = Pt(11)
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except IndexError:
                         pass
                    
    for id, x in zip([0,1,2,3,4,5,6],
                     [1.6,0.3,0.3,0.3,0.3,0.8,0.5]):
        for cell in table.columns[id].cells:
            cell.width = Inches(x)

    table.style = 'Table Grid'

    tablasPath = os.path.join(subareaPath, "Tablas")
    finalPath = os.path.join(tablasPath, "geh_r2_table.docx")
    doc.save(finalPath)

    return finalPath