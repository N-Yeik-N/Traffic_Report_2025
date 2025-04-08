import os
from pathlib import Path
import xml.etree.ElementTree as ET
from unidecode import unidecode
from openpyxl import load_workbook
import pandas as pd
import re

#docx
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

horarios = [
    "00:00 - 05:00",
    "05:00 - 06:30",
    "06:30 - 10:30",
    "10:30 - 12:30",
    "12:30 - 15:00",
    "15:00 - 17:00",
    "17:00 - 22:00",
    "22:00 - 00:00",
    "00:00 - 06:00",
    "06:00 - 12:00",
    "12:00 - 17:00",
    "17:00 - 22:00",
    "22:00 - 00:00",
    ]

def _translation_greens(sigPath: str) -> None:
    tree = ET.parse(sigPath)
    networkTag = tree.getroot()

    #Movements
    sgList = []
    for sg in networkTag.findall("./sgs/sg"):
        sgList.append(int(sg.get("id")))

    #Nombre de las fases
    #sgnames = []
    #for sg in networkTag.findall("./sgs/sg"):
    #    sgnames.append(unidecode(sg.get("name")))
    #print(sgnames)


    #Stages data
    dfStages = pd.DataFrame(
        columns=[int(stage.attrib["id"]) for stage in networkTag.findall('./stages/stage')],
        index=sgList,
    )

    #Filling data of stages and movements
    for stageTag in networkTag.findall("./stages/stage"):
        idStage = int(stageTag.get("id"))
        for activation in stageTag.findall("./activations/activation"):
            index = int(activation.get("sg_id"))
            value = True if activation.get("activation") == "ON" else False
            dfStages.loc[index, idStage] = value
    
    #Finding movements what starts in the beginning
    check = False
    for sg in sgList:
        if dfStages.loc[sg, 1] and not dfStages.loc[sg, dfStages.columns[-1]]: #Only in the first phase
            selected_sg = sg
            check = True
            break

    if not check:
        print(f"Mensaje para el más feo: No hay movimientos que acaben en la última fase y NO inicien en la primera fase")
        print("ERROR: ", sigPath)

    #Finding interstageprog for the selected signal group
    space = 0
    interstageprogTag = networkTag.findall("./interstageProgs/interstageProg")
    for sg in interstageprogTag[-1].findall("./sgs/sg"):
        idSg = int(sg.get("sg_id"))
        if selected_sg == idSg:
            for cmd in sg.findall("./cmds/cmd"):
                if cmd.get("display") == "3":
                    space += int(cmd.get("begin"))//1000
                    break
    
    for stageProg in networkTag.findall("./stageProgs/stageProg"):
        interstage = stageProg.findall("./interstages/interstage")
        if len(interstage) == len(dfStages.columns):
            cycleTime = int(stageProg.attrib["cycletime"])//1000
            lastTime = int(interstage[-1].attrib["begin"])//1000
            upperLimit = cycleTime - space
            movement = upperLimit - lastTime
            for interstage in stageProg.findall("./interstages/interstage"):
                value = interstage.attrib["begin"]
                #Para Yeik: Sub Area 043
                interstage.attrib["begin"] = str(int((int(value)//1000 + movement)*1000))

    ET.indent(tree)
    tree.write(sigPath, encoding="UTF-8", xml_declaration=True)

def _create_from_excel(sig_path, scenarioValue, tipicidadValue, wb):
    codigo = os.path.split(sig_path)[1][:-4]
    ws = wb[codigo]

    #Computing number of phases according peak hours
    listSlicesPeakhours = [
        slice("V4","AJ4"),
        slice("V6","AJ6"),
        slice("V8","AJ8"),
        slice("V10","AJ10"),
        slice("V11","AJ11"),
        slice("V12","AJ12"),
    ]

    maxnumphases = 0
    #emptyThree = lambda three: all(x in [None, ""] for x in three)
    #emptyThree = lambda three: three in [["","",""], [None,None,None], [0,0,0]]
    emptyThree = lambda three: any(x>0 for x in three if isinstance(x, (int, float)))

    for sliceph in listSlicesPeakhours:
        rowList = [elem.value for row in ws[sliceph] for elem in row]
        for i in range(len(rowList)-3, -1, -3):
            terna = rowList[i:i+3]
            if emptyThree(terna):
                numphases = (i+3)//3
                if numphases > maxnumphases:
                    maxnumphases = numphases
                break

    #Extracting data for no peak hours

    listSlices = [
        (slice("V2", "AJ2"), "HPMAD", "Tipico"),
        (slice("V3", "AJ3"), "HVMAD", "Tipico"),
        (slice("V5", "AJ5"), "HVM", "Tipico"),
        (slice("V7", "AJ7"), "HVT", "Tipico"),
        (slice("V9", "AJ9"), "HVN", "Tipico"),  
        (slice("V10", "AJ10"), "HVMAD", "Atipico"),
        (slice("V14", "AJ14"), "HVN", "Atipico"),
    ]

    for slicev, scenario, tipicidad in listSlices:
        if scenarioValue == scenario and tipicidad == tipicidadValue:
            rowList = ['' if elem.value is None else elem.value for row in ws[slicev] for elem in row][:maxnumphases*3]
            greens = [elem for i, elem in enumerate(rowList) if i % 3 == 0]
            ambars = [elem for i, elem in enumerate(rowList) if i % 3 == 1]
            reds = [elem for i, elem in enumerate(rowList) if i % 3 == 2]
            rowListNumbers = [elem for elem in rowList if elem not in ['', None]]
            sig_info = {
                "sig_name": codigo,
                "turn": scenario,
                "tipicidad": tipicidad,
                "cycle_time": sum(rowListNumbers),
                "offset": 0,
                "greens": greens,
                "ambars": ambars,
                "reds": reds
            }
            break
    return sig_info

def _combine_all_docx(filePathMaster, filePathsList, finalPath) -> None:
    number_of_sections = len(filePathsList)
    master = Document(filePathMaster)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document(filePathsList[i])
        composer.append(doc_temp)
    composer.save(finalPath)

def _create_data(sig_path: str, scenario: str, tipicidad: str) -> dict:
    #Traslación de verdes:
    _translation_greens(sig_path)

    tree = ET.parse(sig_path)
    sc_tag = tree.getroot()
    
    #Computing amber and red-red
    greens = []

    #Greens
    maxPhase = (None, 0) #NOTE: Index of stageProg / number of phases
    for i, stageProg in enumerate(sc_tag.findall("./stageProgs/stageProg")):
        interstages = stageProg.findall("./interstages/interstage")
        numberPhases = len(interstages)
        if numberPhases > maxPhase[1]:
            maxPhase = (i, numberPhases)

    stageProg = sc_tag.findall("./stageProgs/stageProg")[maxPhase[0]]
    cycleTime = int(stageProg.get('cycletime'))//1000
    offset = int(stageProg.get('offset'))//1000
    for interstage in stageProg.findall("./interstages/interstage"):
        greens.append(int(interstage.get('begin'))//1000)
        #revisar verdes
        #print(greens)

    firstValue = greens[0]
    greens = [y-x for x,y in zip(greens[:-1], greens[1:])]
    greens[:0] = [firstValue]

    #para revisar
    #print(greens)
    

    ######################
    # Intergreen Matrix  #
    ######################

    intergreenmatrix = sc_tag.find("./intergreenmatrices/intergreenmatrix")

    clearingsg_list = [int(elem_tag.get('clearingsg')) for elem_tag in intergreenmatrix.findall("./intergreen")]
    enteringsg_list = [int(elem_tag.get('enteringsg')) for elem_tag in intergreenmatrix.findall("./intergreen")]
    value_list = [int(elem_tag.get('value'))//1000 for elem_tag in intergreenmatrix.findall("./intergreen")]

    integreen_data = {  
        'clearingsg': clearingsg_list,
        'enteringsg': enteringsg_list,
        'value': value_list,
    }

    df = pd.DataFrame(integreen_data)

    intergreen_matrix = df.pivot(index='clearingsg', columns='enteringsg', values='value')

    #################################
    # Obtaining movements in phases #
    #################################

    # Obtaining number of movements

    movements = [int(sg.get("id")) for sg in sc_tag.findall("./sgs/sg")]

    # Obtaining number of phases

    number_phases = [i+1 for i in range(len(sc_tag.findall("./stages/stage")))]

    # Creating matrix

    df_phases = pd.DataFrame(index=movements, columns=number_phases)

    for nro_phase, stage in enumerate(sc_tag.findall("./stages/stage"), start = 1):
        for activation in stage.findall("./activations/activation"):
            turn_phase = activation.get("activation") == "ON"
            movement = int(activation.get("sg_id"))
            df_phases.loc[movement, nro_phase] = turn_phase

    # Checking if all red exists in a phase

    columns = df_phases.columns
    rr_phases = []
    for i in range(len(columns)):
        col_current = columns[i]
        col_next = columns[(i+1) % len(columns)] # Columna siguiente (cíclica)

        check = False
        for idx, (val_current, val_next) in enumerate(zip(df_phases[col_current], df_phases[col_next])):
            if check:
                break
            # Check if they are T,F or F,T
            if (val_current and not val_next) or (not val_current and val_next):
                # Check if the opposite pair is in the next row
                for next_idx, (next_val_current, next_val_next) in enumerate(zip(df_phases[col_current], df_phases[col_next])):
                    if next_idx != idx and (
                        (not val_current and val_next and next_val_current and not next_val_next) or
                        (val_current and not val_next and not next_val_current and next_val_next)
                    ):
                        rr_phases.append(list(columns)[i])
                        check = True
                        break

    # Obtaining values of ambers
    ambers = []
    for interstageProg in sc_tag.findall("./interstageProgs/interstageProg"):
        checkPhase = False
        for sg in interstageProg.findall("./sgs/sg"):
            fixedState = sg.find("./fixedstates/fixedstate")
            if fixedState is not None and not checkPhase:
                ambers.append(int(fixedState.get("duration"))//1000)
                checkPhase = True
                break
        if not checkPhase: ambers.append(0)

    # Obtaining movements per phase
    movements_by_phase = {}
    for column in list(df_phases.columns):
        movements_by_phase[column] = df_phases.index[df_phases[column]].tolist()

    # Obtaining rr
    reds = []
    for phase, mov_list in movements_by_phase.items():
        if not phase in rr_phases:
            reds.append(0)
            continue
        max_value = 0
        for mov in mov_list:
            #print(intergreen_matrix[mov],"\t",ambers[phase-1])
            #print("____")
            value = int(intergreen_matrix[mov].max())-int(ambers[phase-1]) 
            #print(intergreen_matrix[mov].max(),"\t",ambers[phase-1])
            if value > max_value:
                max_value = value


        reds.append(max_value)

    #corrigiendo verdes
    #restarle a arrelglo greens los ambers y rr de la misma fase
    for i in range(1,len(greens)):
        #print(greens)
        #print(ambers)
        #print(reds,"\n")
        greens[i]=greens[i]-ambers[i-1]-reds[i-1]
        #greens[i]=greens[i]-ambers[i]-reds[i]
        #print(greens,"\n\n\n")
    


    sig_info = {
        "sig_name": os.path.split(sig_path)[1][:-4],
        "turn": scenario,
        "tipicidad": tipicidad,
        "cycle_time": cycleTime,
        "offset": offset,
        "greens": greens,
        "ambars": ambers,
        "reds": reds,
    }

    return sig_info

def _create_table(sigs_info, tablasPath) -> None:
    doc = Document()
    maximum = 0
    for sigInfo in sigs_info:
        valueLen = len(sigInfo['greens'])
        if valueLen > maximum: maximum = valueLen

    len_greens = maximum

    # print("TAMAÑO DE VASES DE VERDE:", len_greens)
    table = doc.add_table(rows = 1, cols= 1+4+len_greens*3)
    table.cell(0,0).text = "Int."
    table.cell(0,1).text = "N° Plan"
    table.cell(0,2).text = "Horario"
    table.cell(0,3).text = "Desfase"
    table.cell(0,4).text = "T.C."
    for i in range(len_greens):
        table.cell(0,4+1+3*i).text = f'Fase {i+1}'
        table.cell(0,4+2+3*i).text = f'A'
        table.cell(0,4+3+3*i).text = f'RR'

    countPlans = 1
    for i, sig_info in enumerate(sigs_info):
        new_row = table.add_row()
        if i == 7:
            atypical_row = table.add_row()
            atypical_row.cells[0].text = "DÍA ATÍPICO"
            atypical_row.cells[0].merge(atypical_row.cells[len(atypical_row.cells)-1])
            
        new_row.cells[1].text = f"{countPlans}" #N° Plan
        new_row.cells[2].text = horarios[countPlans-1] #00:00 - 05:00
        new_row.cells[3].text = f"{sig_info['offset']}" #Desfase
        new_row.cells[4].text = f"{sig_info['cycle_time']}" #Tiempo de Ciclo
        #Repartos:
        for (j, greens), ambars, reds in zip(enumerate(sig_info['greens']),sig_info['ambars'],sig_info['reds']):
            #print(sig_info) aveces falla cuando en el excel program results tiene vacio los valores, deben ir 0
            new_row.cells[4+1+3*j].text = f"{int(greens)}"
            new_row.cells[4+2+3*j].text = f"{int(ambars)}"
            new_row.cells[4+3+3*j].text = f"{int(reds)}"
        countPlans += 1
        
    table.cell(1,0).text = sig_info['sig_name']
    table.cell(1,0).merge(table.cell(8,0))
    table.cell(10,0).text = sig_info['sig_name']
    table.cell(10,0).merge(table.cell(14,0))

    table.style = 'Table Grid'
    table.cell(0,0).width = Cm(1.75)
    table.cell(0,1).width = Cm(1)
    table.cell(0,2).width = Cm(3)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                try:
                    run = paragraph.runs[0]
                except:
                    continue
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(len(table.columns)):
        cell = table.cell(0,i)
        cell_xml_element = cell._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'),'B4C6E7')
        table_cell_properties.append(shade_obj)

    for selectedRow in [0, 9]:
        for i in range(len(table.columns)):
            cell = table.cell(selectedRow,i)
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.bold = True

    finalPath = os.path.join(tablasPath, f"table18_{sig_info['sig_name']}.docx")
    doc.save(finalPath)

    return finalPath, sig_info['sig_name']

def create_table18(subarea_path) -> None:
    #Reading each folder
    tablasPath = os.path.join(subarea_path, "Tablas")
    output_folder = Path(subarea_path) / "Output_Base"
    tipicidades = ["Tipico", "Atipico"]

    scenarioByTipicidad = {
        "Tipico": ["HPMAD", "HVMAD", "HPM", "HVM", "HPT", "HVT", "HPN", "HVN"],
        "Atipico": ["HVMAD", "HPM", "HPT", "HPN", "HVN"],
    }

    #listCodes
    subareaFiles = os.listdir(subarea_path)
    patternInpxFile = r"PTV Vissim Sub Area [0-9]+ \(SA\).inpx"
    vissimFile = [file for file in subareaFiles if re.search(patternInpxFile, file)][0]
    inpxPath = Path(subarea_path) / vissimFile
    tree = ET.parse(inpxPath)
    networkTag = tree.getroot()
    listNodes = []
    for nodeTag in networkTag.findall("./nodes/node"):
        udaTag = nodeTag.find("./uda") #NOTE: There is only one uda in each node
        listNodes.append(udaTag.attrib["value"])
    
    #Program Results:
    programResultsPath = Path(subarea_path) / "Program_Results.xlsx"
    wb = load_workbook(programResultsPath, read_only=True, data_only=True)
    listData = []

    for node in listNodes:
        sigsInfo = []
        for tipicidad in tipicidades:
            for scenario in scenarioByTipicidad[tipicidad]:
                sigPath = output_folder / tipicidad / scenario / f"{node}.sig"
                if scenario in ["HPMAD", "HVMAD", "HVM", "HVT", "HVN"]:
                    sigInfo = _create_from_excel(sigPath, scenario, tipicidad, wb) #NOTE: from program_results.xlsx
                else:
                    sigInfo = _create_data(sigPath, scenario, tipicidad) #NOTE: from .sig
                sigsInfo.append(sigInfo)
        
        finalPath, code = _create_table(sigsInfo, tablasPath)
        texto = f"Programación semafórica de la intersección {node}"
        listData.append((texto, finalPath, node))

    wb.close()

    listWordPaths = []
    for text, pathTable, code in listData:
        doc_template = DocxTemplate("./templates/template_tablas4.docx")
        new_table = doc_template.new_subdoc(pathTable)
        doc_template.render({
            "texto": text,
            "tabla": new_table,
        })
        refPath = os.path.join(
            tablasPath,
            f"table18_{code}_REF.docx"
        )
        doc_template.save(refPath)
        listWordPaths.append(refPath)

    programPath = os.path.join(subarea_path, "Tablas", "Programs.docx")
    filePathMaster = listWordPaths[0]
    filePathList = listWordPaths[1:]
    _combine_all_docx(filePathMaster, filePathList, programPath)

    return programPath

# if __name__ == '__main__':
#     _create_data(
#         sig_path=r"C:\Users\dacan\OneDrive\Desktop\PRUEBAS\Maxima Entropia\03 Proyecto Chorrillos-Barranco\6. Sub Area Vissim\Sub Area 151\Output_Base\Tipico\HPM\VM-35.sig",
#         scenario="HPM",
#         tipicidad="Tipico"
#     )
#     subareaPath = r"D:\Work\02 Proyecto SJL-El Agustino (57 Int. - 18 SA)\6. Sub Area Vissim\Sub Area 080"
#     create_table18(subareaPath)