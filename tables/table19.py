import os
from pathlib import Path
import xml.etree.ElementTree as ET
from unidecode import unidecode


#docx
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

horarios = {
    "Tipico": [
    "00:00 - 05:00",
    "05:00 - 06:30",
    "06:30 - 10:30",
    "10:30 - 12:30",
    "12:30 - 15:00",
    "15:00 - 17:00",
    "17:00 - 22:00",
    "22:00 - 00:00",
    ],
    "Atipico": [
    "00:00 - 06:00",
    "06:00 - 12:00",
    "12:00 - 17:00",
    "17:00 - 22:00",
    "22:00 - 00:00",
    ],
}

def _combine_all_docx(filePathMaster, filePathsList, finalPath) -> None:
    number_of_sections = len(filePathsList)
    master = Document(filePathMaster)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document(filePathsList[i])
        composer.append(doc_temp)

    composer.save(finalPath)

def _create_data(sig_path) -> dict:
    tree = ET.parse(sig_path)
    sc_tag = tree.getroot()
    for stageProg in sc_tag.findall("./stageProgs/stageProg"):
        cycle_time = int(stageProg.attrib['cycletime'])//1000
        offset = int(stageProg.attrib['offset'])//1000
        greens = []
        for interstage in stageProg.findall("./interstages/interstage"):
            greens.append(int(interstage.attrib['begin'])//1000)

    #Mínimo ambars
    min_ambars = []
    min_greens = []
    for sg in sc_tag.findall("./sgs/sg"):
        for defaultDuration in sg.findall("./defaultDurations/defaultDuration"):
            if defaultDuration.attrib['display'] == '4':
                min_ambars.append(int(defaultDuration.attrib['duration'])//1000)
            if defaultDuration.attrib['display'] == '3':
                min_greens.append(int(defaultDuration.attrib['duration'])//1000)

    try:
        min_ambars = min_ambars[:len(greens)]
        min_greens = min_greens[:len(greens)]
    except Exception as inst:
        print("Este error es algo complejo, espero no te salga.")
        raise inst

    cycle_interstages = []
    for interstageProg in sc_tag.findall("./interstageProgs/interstageProg"):
        cycle_interstages.append(int(interstageProg.attrib['cycletime'])//1000)

    min_reds = [x-y-z for x,y,z in zip(cycle_interstages, min_ambars, min_greens)]

    filtered_green = []
    for (i, green), ambar, red in zip(enumerate(greens),min_ambars,min_reds):
        if i == 0: filtered_green.append(green)
        else: filtered_green.append(green - greens[i-1] - ambar - red)

    if sum(filtered_green) + sum(min_ambars) + sum(min_reds) != cycle_time:
        filtered_green[0] = filtered_green[0]+ cycle_time-(sum(filtered_green) + sum(min_ambars) + sum(min_reds))

    sig_info = {
        "sig_name": os.path.split(sig_path)[1][:-4],
        "cycle_time": cycle_time,
        "offset": offset,
        "greens": filtered_green,
        "ambars": min_ambars,
        "reds": min_reds,
    }

    return sig_info

def _create_table(sigs_info, tipicidad, tablasPath) -> None:
    doc = Document()
    sig_info_0 = sigs_info[0]
    greens_0 = sig_info_0['greens']
    len_greens =len(greens_0)

    table = doc.add_table(rows = 1, cols= 1+4+len_greens*3)
    table.cell(0,0).text = "Int."
    table.cell(0,1).text = "N° Plan"
    table.cell(0,2).text = "Horario"
    table.cell(0,3).text = "Desfase"
    table.cell(0,4).text = "T.C."
    for i in range(len_greens):
        table.cell(0,4+1+3*i).text = f'Fase {i+1}'
        table.cell(0,4+2+3*i).text = f'A'
        table.cell(0,4+3+3*i).text = f'RR'

    for i, sig_info in enumerate(sigs_info):
        new_row = table.add_row()
        if i == 0:
            new_row.cells[0].text = sig_info['sig_name'] #Nombre de la intersección
        new_row.cells[1].text = f"{i+1}" #N° Plan
        new_row.cells[2].text = horarios[tipicidad][i] #00:00 - 05:00
        new_row.cells[3].text = f"{sig_info['offset']}" #Desfase
        new_row.cells[4].text = f"{sig_info['cycle_time']}" #Tiempo de Ciclo
        #Repartos:
        for (j, greens), ambars, reds in zip(enumerate(sig_info['greens']),sig_info['ambars'],sig_info['reds']):
            new_row.cells[4+1+3*j].text = f"{greens}"
            new_row.cells[4+2+3*j].text = f"{ambars}"
            new_row.cells[4+3+3*j].text = f"{reds}"
    
    table.cell(1,0).merge(table.cell(i+1,0))

    table.style = 'Table Grid'
    table.cell(0,0).width = Cm(1.75)
    table.cell(0,1).width = Cm(1)
    table.cell(0,2).width = Cm(3)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(len(table.columns)):
        cell = table.cell(0,i)
        cell_xml_element = cell._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'),'B4C6E7')
        table_cell_properties.append(shade_obj)

    for i in range(len(table.columns)):
        cell = table.cell(0,i)
        for paragraph in cell.paragraphs:
            run = paragraph.runs[0]
            run.font.bold = True

    finalPath = os.path.join(tablasPath, f"table18_{sig_info['sig_name']}_{unidecode(tipicidad).upper()}.docx")
    doc.save(finalPath)
    return finalPath, sig_info['sig_name'], tipicidad

def create_table19(subarea_path) -> None:
    #Reading each folder
    tablasPath = os.path.join(subarea_path, "Tablas")
    output_folder = Path(subarea_path) / "Output_Proyectado"
    tipicidades = ["Tipico", "Atipico"]

    listData = []
    for tipicidad in tipicidades:
        scenarios = os.listdir(output_folder / tipicidad)
        
        for i, scenario in enumerate(scenarios):
            if i == 0:
                scenario_path = output_folder / tipicidad / scenario
                sig_files = os.listdir(scenario_path)
                sig_files = [file for file in sig_files if file.endswith(".sig")]
            else:
                break

        for sig_file in sig_files:
            sigs_info = []
            for scenario in scenarios:
                sig_path = output_folder / tipicidad / scenario / sig_file
                sig_info = _create_data(sig_path) #Data necesaria por fases de una intersección.
                sigs_info.append(sig_info) #Data por cada horario de los programas (HVMAD, HPM, ...)
     
            finalPath, code, tipicidad = _create_table(sigs_info, tipicidad, tablasPath)
            texto = f"Programación semafórica de la intersección {code} día {tipicidad}"
            listData.append((texto, finalPath, code, tipicidad)) 

    listWordPaths = []
    for text, pathTable, code, tipicidad in listData:
        doc_template = DocxTemplate("./templates/template_tablas4.docx")
        new_table = doc_template.new_subdoc(pathTable)
        doc_template.render({
            "texto": text,
            "tabla": new_table,
        })
        refPath = os.path.join(
            tablasPath,
            f"table18_1Y_{code}_{unidecode(tipicidad).upper()}_REF.docx"
        )
        doc_template.save(refPath)
        listWordPaths.append(refPath)

    programPath = os.path.join(subarea_path, "Tablas", "Programs1Y.docx")
    filePathMaster = listWordPaths[0]
    filePathList = listWordPaths[1:]
    _combine_all_docx(filePathMaster, filePathList, programPath)

    return programPath

# if __name__ == '__main__':
#     path = r"C:\Users\dacan\OneDrive\Desktop\PRUEBAS\Maxima Entropia\04 Proyecto Universitaria (37 Int. - 19 SA)\6. Sub Area Vissim\Sub Area 016"
#     read_programs(path)