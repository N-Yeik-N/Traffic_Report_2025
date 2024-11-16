import os
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import os
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

tipicoList = ["HPM", "HPT", "HPN"]
atipicoList = ["HPM", "HPT", "HPN"]

def _remove_top_bottom_bordes(cell, topValue = True, bottomValue = True):
    tc_pr = cell._element.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    if topValue:
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'nil')
        tc_borders.append(top)
    
    if bottomValue:
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'nil')
        tc_borders.append(bottom)

    tc_pr.append(tc_borders)

def get_color_by_los(los):
    colores = {
        "A": "00B050", # Verde
        "B": "B5E6A2", # Verde amarillento
        "C": "FFFF99", # Amarillo
        "D": "FFD961", # Naranja
        "E": "EB844B", # Naranja rojizo
        "F": "FF3B3B", # Rojo
    }
    return colores.get(los, "FFFFFF") #Blanco por defecto

def _align_content(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial Narrow'

    for i in range(len(table.columns)):
        cell = table.cell(0,i)
        cell_xml_element = cell._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'),'B4C6E7')
        table_cell_properties.append(shade_obj)

    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            try:
                run = paragraph.runs[0]
                run.font.bold = True
            except IndexError:
                pass

    for row in table.rows:
        for col_index in range(3):  # Columnas 0, 1 y 2
            cell = row.cells[col_index]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def create_table23(subareaPath):
    actualPath = os.path.join(subareaPath, "Actual")
    basePath = os.path.join(subareaPath, "Output_Base")
    proyectadoPath = os.path.join(subareaPath, "Output_Proyectado")

    listJSONPathsActual = []
    listJSONPathsBase = []
    listJSONPathProyectado = []

    listNames = []
    for tipicidad in ["Tipico", "Atipico"]:
        #Actual
        tipicidadPathActual = os.path.join(actualPath, tipicidad)
        scenariosListActual = os.listdir(tipicidadPathActual)
        scenariosListActual = [file for file in scenariosListActual if not file.endswith(".ini") and file in ["HPM", "HPT", "HPN"]]

        #Output Base
        tipicidadPathBase = os.path.join(basePath, tipicidad)
        scenariosListBase = os.listdir(tipicidadPathBase)
        scenariosListBase = [file for file in scenariosListBase if not file.endswith(".ini") and file in ["HPM", "HPT", "HPN"]]

        #Output Proyectado
        tipicidadPathProyectado = os.path.join(proyectadoPath, tipicidad)
        scenariosListProyectado = os.listdir(tipicidadPathProyectado)
        scenariosListProyectado = [file for file in scenariosListProyectado if not file.endswith(".ini") and file in ["HPM", "HPT", "HPN"]]

        if tipicidad == "Tipico":
            for tipicoUnit in tipicoList:
                for i in range(len(scenariosListActual)):
                    if tipicoUnit == scenariosListActual[i]:

                        scenarioPathActual = os.path.join(tipicidadPathActual, scenariosListActual[i])
                        scenarioContentActual = os.listdir(scenarioPathActual)
                        if "table.json" in scenarioContentActual:
                            jsonFileActual = os.path.join(scenarioPathActual, "table.json")
                            listJSONPathsActual.append(jsonFileActual)    
                            listNames.append(scenariosListActual[i])

                        scenarioPathBase = os.path.join(tipicidadPathBase, scenariosListBase[i])
                        scenarioContentBase = os.listdir(scenarioPathBase)
                        if "table.json" in scenarioContentBase:
                            jsonFileBase = os.path.join(scenarioPathBase, "table.json")
                            listJSONPathsBase.append(jsonFileBase)

                        scenarioPathProyectado = os.path.join(tipicidadPathProyectado, scenariosListProyectado[i])
                        scenarioContentProyectado = os.listdir(scenarioPathProyectado)
                        if "table.json" in scenarioContentProyectado:
                            jsonFileProyectado = os.path.join(scenarioPathProyectado, "table.json")
                            listJSONPathProyectado.append(jsonFileProyectado)

        elif tipicidad == "Atipico":
            for tipicoUnit in tipicoList:
                for i in range(len(scenariosListActual)):
                    if tipicoUnit == scenariosListActual[i]:
                        
                        scenarioPathActual = os.path.join(tipicidadPathActual, scenariosListActual[i])
                        scenarioContentActual = os.listdir(scenarioPathActual)
                        if "table.json" in scenarioContentActual:
                            jsonFileActual = os.path.join(scenarioPathActual, "table.json")
                            listJSONPathsActual.append(jsonFileActual)    
                            listNames.append(scenariosListActual[i])

                        scenarioPathBase = os.path.join(tipicidadPathBase, scenariosListBase[i])
                        scenarioContentBase = os.listdir(scenarioPathBase)
                        if "table.json" in scenarioContentBase:
                            jsonFileBase = os.path.join(scenarioPathBase, "table.json")
                            listJSONPathsBase.append(jsonFileBase)

                        scenarioPathProyectado = os.path.join(tipicidadPathProyectado, scenariosListProyectado[i])
                        scenarioContentProyectado = os.listdir(scenarioPathProyectado)
                        if "table.json" in scenarioContentProyectado:
                            jsonFileProyectado = os.path.join(scenarioPathProyectado, "table.json")
                            listJSONPathProyectado.append(jsonFileProyectado)

    #####################
    # Vehicular results #
    #####################

    doc = Document()
    table = doc.add_table(rows = 1, cols = 9)
    table.style = 'Table Grid'

    table.cell(0,0).text = "Tipicidad"
    table.cell(0,1).text = "Escenario"
    table.cell(0,1).merge(table.cell(0,2))

    #Creating Headers
    for i, texto in enumerate([
        "Nodo", "Número de\nVehículos\n(veh)", "Cola Máx.\nPromedio\n(m)", "Demora por parada\nPromedio\n(s/veh)", "Demora\nPromedio\n(s/veh)", "LOS\n(A-F)"
        ]):
        table.cell(0,i+3).text = texto

    nroRow = 1

    for i, jsonPathActual in enumerate(listJSONPathsActual):
        with open(jsonPathActual, 'r') as file:
            data = json.load(file)

        try:
            with open(listJSONPathsBase[i], 'r') as file2:
                data2 = json.load(file2)
            checkBase = True
        except (FileNotFoundError, IndexError) as e:
            checkBase = False

        try:
            with open(listJSONPathProyectado[i], 'r') as file3:
                data3 = json.load(file3)
            checkProyectado = True
        except (FileNotFoundError, IndexError) as e:
            checkProyectado = False

        for j in range(len(data['nodes_names'])): #NOTE: Estoy considerando que son del mismo tamaño, puede que no sea así siempre.
            
            #Actual
            new_row = table.add_row()
            new_row.cells[2].text = "Actual"
            _remove_top_bottom_bordes(new_row.cells[2], topValue=False)
            new_row.cells[3].text = data['nodes_names'][j]                  #Nodo
            new_row.cells[4].text = str(round(float(data['nodes_totres'][j][4])))    #Número de Vehículos
            new_row.cells[5].text = str(round(float(data['nodes_totres'][j][3])))    #Cola Máx. Promedio
            new_row.cells[6].text = str(round(float(data['nodes_totres'][j][1]), 1))    #Pare Promedio
            new_row.cells[7].text = str(round(float(data['nodes_totres'][j][0]), 1))    #Demora Promedio
            new_row.cells[8].text = data['nodes_los'][j]
            for selectedColumns in [3, 4, 5, 6, 7, 8]:
                _remove_top_bottom_bordes(new_row.cells[selectedColumns], topValue=False)
            color_hex = get_color_by_los(data['nodes_los'][j])
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
            new_row.cells[8]._element.get_or_add_tcPr().append(shading_elm)
            nroRow += 1

            #Base
            if checkBase:
                new_row = table.add_row()
                table.cell(nroRow-1,3).merge(table.cell(nroRow,3))
                new_row.cells[2].text = "Propuesto"
                _remove_top_bottom_bordes(new_row.cells[2])
                new_row.cells[4].text = str(round(float(data2['nodes_totres'][j][4])))    #Número de Vehículos
                new_row.cells[5].text = str(round(float(data2['nodes_totres'][j][3])))    #Cola Máx. Promedio
                new_row.cells[6].text = str(round(float(data2['nodes_totres'][j][1]), 1))    #Pare Promedio
                new_row.cells[7].text = str(round(float(data2['nodes_totres'][j][0]), 1))    #Demora Promedio
                new_row.cells[8].text = data2['nodes_los'][j]
                for selectedColumns in [4, 5, 6, 7, 8]:
                    _remove_top_bottom_bordes(new_row.cells[selectedColumns])
                color_hex = get_color_by_los(data2['nodes_los'][j])
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
                new_row.cells[8]._element.get_or_add_tcPr().append(shading_elm)
                nroRow += 1

            #Proyectado
            if checkProyectado:
                new_row = table.add_row()
                table.cell(nroRow-1,3).merge(table.cell(nroRow,3))
                #table.cell(nroRow-1,2).merge(table.cell(nroRow,2))
                new_row.cells[2].text = "Proyectado"
                _remove_top_bottom_bordes(new_row.cells[2], bottomValue=False)
                new_row.cells[4].text = str(round(float(data3['nodes_totres'][j][4])))    #Número de Vehículos
                new_row.cells[5].text = str(round(float(data3['nodes_totres'][j][3])))    #Cola Máx. Promedio
                new_row.cells[6].text = str(round(float(data3['nodes_totres'][j][1]), 1))    #Pare Promedio
                new_row.cells[7].text = str(round(float(data3['nodes_totres'][j][0]), 1))    #Demora Promedio
                new_row.cells[8].text = data3['nodes_los'][j]
                for selectedColumns in [4, 5, 6, 7, 8]:
                    _remove_top_bottom_bordes(new_row.cells[selectedColumns], bottomValue=False)
                color_hex = get_color_by_los(data3['nodes_los'][j])
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
                new_row.cells[8]._element.get_or_add_tcPr().append(shading_elm)
                nroRow += 1

        if i == 0:
            numberNodes = len(data['nodes_names'])

    numberTurns = 1 #Actual +1
    if checkBase:
        numberTurns += 1
    if checkProyectado:
        numberTurns += 1
    JUMP = numberNodes*numberTurns-1

    #(ni.ns-1) == JUMP
    startRow = 1
    for turnName in ['HPM', 'HPT', 'HPN']*2:
        table.cell(startRow, 1).text = turnName
        table.cell(startRow, 1).merge(table.cell(startRow+JUMP, 1))
        startRow = startRow+JUMP+1
    
    #Tipicidad
    totalRows = len(table.rows) - 1
    table.cell(1,0).text = "TÍPICO"
    table.cell(1,0).merge(table.cell(totalRows//2,0))
    table.cell(totalRows//2+1,0).text = "ATÍPICO"
    table.cell(totalRows//2+1,0).merge(table.cell(totalRows,0))

    _align_content(table)

    finalPathVehicle = os.path.join(subareaPath, "Tablas", "resumenTable.docx")
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(finalPathVehicle)

    ######################
    # Pedestrian results #
    ######################

    doc = Document()
    table = doc.add_table(rows = 19, cols=5)
    table.style = 'Table Grid'

    #Headers
    table.cell(0,0).text = "Tipicidad"
    table.cell(0,1).text = "Escenario"
    table.cell(0,1).merge(table.cell(0,2))
    table.cell(0,3).text = "Vel. Prom.\n(km/h)"
    table.cell(0,4).text = "Tiempo Viaje Prom.\n(seg/peat)"

    #Tipically and turn
    table.cell(1,0).text = "TÍPICO"
    table.cell(1,0).merge(table.cell(9,0))

    table.cell(10,0).text = "ATÍPICO"
    table.cell(10,0).merge(table.cell(18,0))

    for i in range(2):
        table.cell(1+9*i,1).text = "HPM"
        table.cell(1+9*i,1).merge(table.cell(1+9*i+2,1))
        
        table.cell(4+9*i,1).text = "HPT"
        table.cell(4+9*i,1).merge(table.cell(4+9*i+2,1))

        table.cell(7+9*i,1).text = "HPN"
        table.cell(7+9*i,1).merge(table.cell(7+9*i+2,1))

    #Labels per scenario
    for i in range(6):
        table.cell(1+3*i,2).text = "Actual"
        table.cell(2+3*i,2).text = "Propuesto"
        table.cell(3+3*i,2).text = "Proyectado"
        for col in [2,3,4]:
            _remove_top_bottom_bordes(table.cell(1+3*i,col), topValue=False)
            _remove_top_bottom_bordes(table.cell(2+3*i,col))
            _remove_top_bottom_bordes(table.cell(3+3*i,col), bottomValue=False)

    jump = 0
    for tipicidad in ["Tipico", "Atipico"]:
        for turno in ["HPM", "HPT", "HPN"]:
            #Paths
            actualPath = os.path.join(
                subareaPath, "Actual", tipicidad, turno, 'table.json'
            )
            basePath = os.path.join(
                subareaPath, "Output_Base", tipicidad, turno, 'table.json'
            )
            proyectadoPath = os.path.join(
                subareaPath, "Output_Proyectado", tipicidad, turno, 'table.json'
            )

            with open(actualPath, 'r') as file1:
                data1 = json.load(file1)
            with open(basePath, 'r') as file2:
                data2 = json.load(file2)
            with open(proyectadoPath, 'r') as file3:
                data3 = json.load(file3)

            for i in range(3):
                table.cell(1+3*jump,3).text = f"{float(data1['pedestrian_performance']['Avg']['SpeedAvg']):.2f}"
                table.cell(1+3*jump,4).text = str(int(float(data1['pedestrian_performance']['Avg']['TravTmAvg'])))
                table.cell(2+3*jump,3).text = f"{float(data2['pedestrian_performance']['Avg']['SpeedAvg']):.2f}"
                table.cell(2+3*jump,4).text = str(int(float(data2['pedestrian_performance']['Avg']['TravTmAvg'])))
                table.cell(3+3*jump,3).text = f"{float(data3['pedestrian_performance']['Avg']['SpeedAvg']):.2f}"
                table.cell(3+3*jump,4).text = str(int(float(data3['pedestrian_performance']['Avg']['TravTmAvg'])))
            jump += 1

    _align_content(table)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in table.rows:
        for cell in row.cells[:3]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    finalPathPedestrian = os.path.join(
        subareaPath, "Tablas", "Results", "resumenPedestrianTable.docx"
    )
    doc.save(finalPathPedestrian)

    return finalPathVehicle, finalPathPedestrian